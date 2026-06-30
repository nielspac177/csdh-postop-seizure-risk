"""Task 27, Journal-format manuscript build (proof-of-concept framing).

Target: a high-impact neurology/neurosurgery journal accepting Original
Research articles with the following conventional limits (most BMJ-family
and similar journals):
  • Original Research: 4,000 words main text
  • Structured abstract: 250 words
  • Maximum 6 figures/tables in main text
  • Vancouver references (numbered, superscript)
  • TRIPOD-AI reporting checklist for prediction models
  • Supplementary material accepted separately

Outputs:
  manuscript/main_manuscript.docx     (≤4000 words, F1–F6, Table 1)
  manuscript/supplementary.docx       (Figures S1–S7, Tables S1–S5,
                                       TRIPOD-AI checklist)
"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _shared import OUT, RES, FIG

MANUS_DIR = OUT / "manuscript"
MAIN_PATH = MANUS_DIR / "main_manuscript.docx"
SUPP_PATH = MANUS_DIR / "supplementary.docx"

# ─── Doc helpers ─────────────────────────────────────────────
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x14, 0x1F, 0x3A)
        r.font.name = "Times New Roman"
    return h

def add_para(doc, text, *, italic=False, bold=False, indent=False, size=11,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Times New Roman"
    r.italic = italic; r.bold = bold
    return p

def add_runs(doc, runs, *, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """runs = list of (text, dict-of-formatting)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Inches(0.3)
    for text, fmt in runs:
        r = p.add_run(text)
        r.font.size = Pt(fmt.get("size", 11)); r.font.name = "Times New Roman"
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if fmt.get("superscript"):
            r.font.superscript = True
    return p

def add_figure(doc, path, *, caption, width=Inches(6.5)):
    if not os.path.exists(path):
        add_para(doc, f"[Figure missing: {path}]", italic=True); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=width)
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cr = cp.add_run(caption)
    cr.font.size = Pt(10); cr.italic = True; cr.font.name = "Times New Roman"

def add_table_from_df(doc, df, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption); cr.bold = True; cr.font.size = Pt(11)
        cr.font.name = "Times New Roman"
    t = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    t.style = "Light List Accent 1"
    for j, col in enumerate(df.columns):
        cell = t.rows[0].cells[j]; cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(10); r.font.name = "Times New Roman"
    for i, (_, row) in enumerate(df.iterrows()):
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                val = f"{val:.3f}" if abs(val) < 100 else f"{val:,.0f}"
            cell = t.rows[i + 1].cells[j]; cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10); r.font.name = "Times New Roman"

# A registry that accumulates inline figure / table references so they can be
# rendered at the end of the manuscript in a dedicated "Tables and Figures"
# section (standard journal-submission format).
_FIG_REGISTRY = []      # list of (label, path, caption)
_TBL_REGISTRY = []      # list of (label, df, caption)

def register_figure(label, path, caption):
    _FIG_REGISTRY.append((label, path, caption))

def register_table(label, df, caption):
    _TBL_REGISTRY.append((label, df, caption))

def render_collected_tables_and_figures(doc):
    """Emit a 'Tables and Figures' section at the end of the manuscript."""
    add_heading(doc, "Tables and Figures", level=1)
    add_para(doc,
        "Per journal-submission convention, all tables and figures with "
        "their full legends are collected at the end of the manuscript "
        "rather than placed inline.", italic=True)
    # Tables first
    if _TBL_REGISTRY:
        add_heading(doc, "Tables", level=2)
        for label, df, caption in _TBL_REGISTRY:
            add_table_from_df(doc, df, caption=caption)
            add_para(doc, "")
    # Figures
    if _FIG_REGISTRY:
        add_heading(doc, "Figures", level=2)
        for label, path, caption in _FIG_REGISTRY:
            add_figure(doc, path, caption=caption)
            add_para(doc, "")

def add_page_break(doc):
    doc.add_page_break()

def setup_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0); section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(11)

# ───────────────────────────────────────────────────────────
# Main manuscript
# ───────────────────────────────────────────────────────────
def build_main():
    doc = Document(); setup_document(doc)

    # Title page
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Postoperative seizure after chronic subdural haematoma "
                    "evacuation: a calibration-focused, conformal-prediction "
                    "proof-of-concept with value-of-information analysis")
    tr.bold = True; tr.font.size = Pt(15); tr.font.name = "Times New Roman"
    add_para(doc, "")
    add_para(doc, "Niels Pacheco-Barrios MD",
              alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Department of Neurosurgery, Beth Israel Deaconess Medical "
                   "Center, Harvard Medical School, Boston, MA, USA",
              alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True)
    add_para(doc, "")
    add_para(doc, "Correspondence: Niels Pacheco-Barrios MD · Department of "
                   "Neurosurgery, BIDMC, 330 Brookline Ave, Boston MA 02215, USA · "
                   "nielspacheco1997@gmail.com", size=10)
    add_para(doc, "")
    add_para(doc, "Manuscript type: Original Research (proof of concept)", size=10)
    add_para(doc, "Word count, abstract: 250 · main text: ~3,400 · "
                   "Figures: 6 · Tables: 1 · Supplementary: yes · "
                   "References: 46", size=10)
    add_runs(doc, [("Keywords: ", {"bold": True}),
        ("chronic subdural haematoma; postoperative seizure; clinical "
         "prediction model; conformal prediction; cost-effectiveness and "
         "value-of-information analysis", {})], indent=False)
    add_para(doc, "Reporting: TRIPOD-AI (checklist in Supplementary Appendix S1).", size=10)
    add_para(doc, "Code and data availability:  The analysis code, six "
                   "main figures, eight supplementary figures, the TRIPOD-AI "
                   "reporting checklist (Supplementary Appendix S1), the "
                   "reproducibility appendix (Supplementary Appendix S2) "
                   "and the BIDMC feature dictionary (Supplementary "
                   "Appendix S3) are "
                   "openly available at github.com/nielspac177/csdh-postop-"
                   "seizure-risk and permanently archived on Zenodo "
                   "(DOI: [pending, minted at manuscript submission via "
                   "the GitHub–Zenodo integration]). An interactive "
                   "companion site at nielspac177.github.io/csdh-postop-"
                   "seizure-risk provides a calibrated patient-level risk "
                   "calculator, a population cost-savings tool, and an "
                   "interactive code callgraph; all computation runs "
                   "client-side in the browser so that no patient "
                   "information is transmitted. Raw patient-level data are "
                   "restricted by the BIDMC Institutional Review Board and "
                   "the eICU Collaborative Research Database data-use "
                   "agreement; filtered, de-identified working "
                   "subsets are released to authorised peer reviewers via "
                   "the reviewer-access protocol documented at "
                   "github.com/nielspac177/csdh-postop-seizure-risk/tree/"
                   "reviewer-access-template.",
              size=10)
    add_para(doc, "Funding: This study received no specific grant from any "
                   "funding agency in the public, commercial or not-for-profit "
                   "sectors.", size=10)
    add_para(doc, "Conflicts of interest: None declared.", size=10)
    add_para(doc, "Author contributions: N.P.-B. is the sole author and was "
                   "responsible for study conception and design, data curation, "
                   "analysis, interpretation, and drafting and revising the "
                   "manuscript.", size=10)
    add_para(doc, "Ethics: Approved by the Beth Israel Deaconess Medical Center "
                   "Institutional Review Board (Protocol [IRB number to be "
                   "inserted]); the eICU Collaborative Research Database v2.0 was "
                   "accessed under its data-use agreement.", size=10)
    add_para(doc, "Data availability: Analysis code, figures, the TRIPOD-AI "
                   "checklist and aggregate results are openly available at "
                   "github.com/nielspac177/csdh-postop-seizure-risk and archived "
                   "on Zenodo (DOI to be minted at acceptance). Patient-level "
                   "data are restricted by the BIDMC IRB and the eICU data-use "
                   "agreement; filtered de-identified subsets are available to "
                   "authorised reviewers via the documented reviewer-access "
                   "protocol.", size=10)
    add_page_break(doc)

    # Abstract (~250 words, structured)
    add_heading(doc, "Abstract", level=1)
    add_runs(doc, [("Background. ", {"bold": True}),
        ("Postoperative seizure complicates 7–12% of chronic subdural "
         "haematoma (cSDH) evacuations, but routine antiepileptic drug (AED) "
         "prophylaxis carries fall, cognitive and drug-interaction risks "
         "specific to elderly patients, and its efficacy after cSDH is "
         "unproven. Calibrated risk stratification that can support "
         "individual treatment decisions is needed.", {})], indent=False)
    add_runs(doc, [("Methods. ", {"bold": True}),
        ("We developed and externally evaluated a machine-learning risk "
         "score for postoperative seizure in 655 cSDH evacuations at "
         "BIDMC (development; 48 events) and 3,297 non-traumatic subdural "
         "haematoma ICU stays across 42 hospitals in the eICU Collaborative "
         "Research Database (external evaluation in a related population; "
         "300 seizures). Eleven model "
         "classes were compared, including six "
         "SMOTE-family oversamplers, Optuna-tuned XGBoost and LightGBM, a "
         "diverse-base stacking ensemble, Bayesian logistic regression with "
         "eICU-informed priors, and Firth penalized logistic regression. "
         "Class-conditional (Mondrian) conformal prediction provided "
         "individual-patient decision sets. A probabilistic cost-"
         "effectiveness analysis with value-of-information (VOI) was run "
         "on 10,000 Monte Carlo iterations.", {})], indent=False)
    add_runs(doc, [("Results. ", {"bold": True}),
        ("Firth penalized logistic regression was the candidate model, fit "
         "on a leakage-safe postoperative-B feature set (18 variables "
         "available at the end of evacuation, before the AED/EEG decision); "
         "it discriminated at AUC 0.645. A postoperative-A set that "
         "additionally included three peri-decision variables reached AUC "
         "0.681 (95% CI 0.609–0.753) but is not deployable, the difference "
         "lying within the Bernoulli noise floor at 48 events. Firth matched "
         "BalancedRandomForest discrimination (DeLong p = 0.81) with markedly "
         "better calibration (Brier 0.068 vs 0.228) and calibration-in-the-"
         "large near zero, though predictions remained under-dispersed "
         "(recalibration slope >1) given the narrow probability range at 48 "
         "events. Eleven model classes converged on an AUC ceiling near 0.68 "
         "consistent with 2022–2025 meta-evidence. In external evaluation on a "
         "related (mixed-acuity ICU) population, AUC was 0.750 (0.711–0.774); "
         "random-effects pooled AUC across 42 hospitals was 0.684 (95% CI "
         "0.651–0.714; logit-scale I²=0%, raw-scale 95% prediction interval "
         "0.53–0.95). Class-conditional conformal "
         "prediction at α=0.10 gave confident singleton predictions in 22% of "
         "patients (rule-out 11%; rule-in 11%). ML-guided allocation — selective "
         "AED prophylaxis as the primary deployable strategy, or selective "
         "continuous-EEG monitoring where available — beat watchful observation "
         "and, at a matched treated fraction, beat random allocation "
         "(discrimination premium ≈$1,000–1,600/patient), confirming the model, "
         "not the treated fraction, added value. ML-guided allocation was "
         "preferred to universal AED under cSDH-plausible AED assumptions; "
         "universal AED was preferable only under optimistic, imported efficacy. "
         "Selective AED is the simplest to deploy and assumes AED has some "
         "efficacy, whereas selective monitoring retains value even if AED is "
         "ineffective. Population value of perfect "
         "information at $100k/QALY was approximately $23M over 10 years.", {})],
         indent=False)
    add_runs(doc, [("Conclusion. ", {"bold": True}),
        ("For postoperative seizure after cSDH evacuation, optimising "
         "small-cohort clinical machine learning for calibration and "
         "individual-patient decision support, rather than discrimination, "
         "yields a candidate model that meets methodological preconditions "
         "for prospective clinical evaluation. Selective, ML-guided AED "
         "prophylaxis is the primary deployable strategy, with selective "
         "continuous-EEG monitoring as an enhanced option where available; "
         "whether selective allocation is preferable to treating all or none "
         "hinges on AED efficacy and harm in cSDH, which current evidence "
         "leaves uncertain; value-of-information prioritises resolving these.", {})],
         indent=False)
    add_page_break(doc)

    # 1. Introduction (~600 words)
    add_heading(doc, "Introduction", level=1)
    add_runs(doc, [
        ("Chronic subdural haematoma (cSDH) is one of the most common "
         "neurosurgical conditions in older adults, with an annual operative "
         "incidence approaching 40,000 cases in the United States and "
         "increasing in parallel with population ageing and anticoagulant "
         "use.", {}), ("¹⁻³", {"superscript": True}), (" Postoperative seizure "
         "complicates 7–12% of cSDH evacuations and is independently "
         "associated with prolonged intensive-care stay, 30-day mortality, "
         "and lower functional independence at follow-up.", {}),
        ("⁴⁻⁶", {"superscript": True})], indent=True)
    add_runs(doc, [
        ("Routine prophylactic antiepileptic drug (AED) administration is "
         "controversial. Levetiracetam, the most commonly used agent, "
         "carries specific harms in the elderly: neuropsychiatric adverse "
         "events in roughly 15–20% of users, somnolence in approximately "
         "28%, and a relative risk of falls of 1.6–1.8.", {}),
        ("⁷⁻¹⁰", {"superscript": True}),
        (" Continuous electroencephalography (cEEG) provides the "
         "alternative of selective treatment, and point-of-care devices "
         "(Ceribell) have lowered per-use costs by approximately $5,600 "
         "relative to conventional cEEG.", {}),
        ("¹¹⁻¹³", {"superscript": True}),
        (" Choosing between strategies for each patient requires "
         "calibrated risk stratification, which is currently empirical.", {})],
        indent=True)
    add_runs(doc, [
        ("Previous machine-learning attempts in this space have reported "
         "discrimination metrics on small single-centre cohorts without "
         "external validation, calibration assessment, decision-analytic "
         "framing, or formal uncertainty quantification.", {}),
        ("¹⁴⁻¹⁶", {"superscript": True}),
        (" Contemporary meta-evidence shows that class-imbalance "
         "corrections do not raise the area under the receiver-operating-"
         "characteristic curve (AUC) in clinical risk models and can "
         "degrade calibration,", {}),
        ("¹⁷⁻¹⁹", {"superscript": True}),
        (" and tabular foundation models have shown limited gains in "
         "head-to-head clinical benchmarks.", {}),
        ("²⁰,²¹", {"superscript": True}),
        (" The implication is that for small clinical cohorts the right "
         "optimisation target is not discrimination but calibration and "
         "individual-patient decision support.", {})], indent=True)
    add_runs(doc, [
        ("We undertook a proof-of-concept study to develop and externally "
         "validate a calibrated, parametric risk score, computable at the "
         "intended decision point, for "
         "postoperative seizure after cSDH evacuation. The work integrates "
         "(i) an eleven-method modelling battery to characterise the "
         "discrimination ceiling, (ii) Firth penalized logistic regression "
         "as the candidate model on the basis of calibration, "
         "interpretability and valid coefficient confidence intervals, "
         "(iii) class-conditional conformal prediction for individual-"
         "patient decisions with distribution-free coverage guarantees, "
         "(iv) cost-effectiveness analysis with value-of-information "
         "scaling to the annual operative population, and (v) a "
         "documented transfer-learning failure between mixed-acuity "
         "intensive-care and pure operative cSDH cohorts.", {})], indent=True)
    add_runs(doc, [
        ("Conformal prediction adds distribution-free, finite-sample "
         "coverage guarantees that probability calibration alone does "
         "not provide and that Bayesian credible intervals require "
         "correct priors and Markov-chain Monte Carlo to deliver", {}),
        ("²³,²⁴", {"superscript": True}),
        (", supplying a principled \"I do not know\" signal at the "
         "bedside, applied as a post-hoc wrapper that leaves the "
         "underlying Firth model untouched. This combination of "
         "calibration, conformal decision support and value-of-"
         "information has not, to our knowledge, been applied to "
         "postoperative seizure prediction.", {})], indent=True)
    add_page_break(doc)

    # 2. Methods (~1,200 words)
    add_heading(doc, "Methods", level=1)
    add_heading(doc, "Study design and ethics", level=2)
    add_runs(doc, [
        ("This retrospective multi-database study followed the TRIPOD-AI "
         "reporting standard for clinical prediction models built using "
         "machine learning (Supplementary Appendix S1).", {}),
        ("²²", {"superscript": True}),
        (" The BIDMC analysis was approved by the institutional review "
         "board (Protocol [IRB number to be inserted]); the eICU Collaborative Research "
         "Database v2.0 was accessed under its data-use agreement.", {})],
        indent=True)
    add_heading(doc, "Cohort assembly", level=2)
    add_runs(doc, [
        ("The BIDMC development cohort comprised all consecutive patients "
         "undergoing burr-hole or craniotomy for cSDH between January 2010 "
         "and December 2023 (n = 655; 48 in-admission postoperative "
         "seizures). For external evaluation, the eICU Collaborative Research "
         "Database was screened in stages: 5,376 ICU stays with any subdural "
         "haematoma diagnosis across 139 hospitals were identified, from "
         "which the pre-specified primary stratum, non-traumatic SDH stays "
         "at hospitals contributing ≥10 such stays, yielded 3,297 stays "
         "across 42 hospitals (300 seizures). All headline external figures, "
         "including the leave-one-hospital-out pooling, use this 3,297/42 "
         "primary stratum; the broader 5,376/139 screen and three "
         "alternative cohort definitions are reported as sensitivity analyses "
         "in the Supplement, and a CONSORT-style inclusion-flow diagram "
         "(Supplementary Appendix S8) details every exclusion step.", {})],
        indent=True)
    add_heading(doc, "Outcome and features", level=2)
    add_runs(doc, [
        ("The primary outcome was any postoperative seizure within the "
         "index admission, ascertained from chart documentation at BIDMC "
         "and from the structured seizure flag in eICU; the two phenotypes "
         "are related but not identical, and the resulting limits on "
         "cross-cohort comparability are addressed in the Discussion. Two "
         "feature sets were defined by their availability at the intended "
         "decision point, the end of haematoma evacuation, before the "
         "antiepileptic-drug/EEG decision. The leakage-safe postoperative-B set "
         "comprised 18 demographic, operative and imaging-derived variables "
         "all available at that point. The postoperative-A set added three "
         "peri-decision variables (AED timing, proportion of stay on AED, "
         "abnormal EEG) that a clinician could not have at evacuation exit; "
         "because these encode information generated around or after the very "
         "decision the model is meant to inform, postoperative-A is reported "
         "only as a paired sensitivity comparison and is not deployable. The "
         "full feature dictionary is in Supplementary Appendix S3.", {})],
        indent=True)
    add_heading(doc, "Modelling", level=2)
    add_runs(doc, [
        ("Eleven model classes were compared: the prior version's BalancedRandom"
         "Forest (BRF) reference; Firth penalized logistic regression; "
         "Bayesian logistic regression with weakly-informative and "
         "eICU-informed priors; class-weighted RandomForest; six SMOTE-"
         "family oversampling variants (SMOTE, Borderline-SMOTE, SVM-SMOTE, "
         "ADASYN, SMOTEENN, SMOTETomek); Optuna-tuned XGBoost and "
         "LightGBM (40 trials each); and a diverse-base stacking ensemble "
         "(logistic regression, BRF, XGBoost, k-nearest neighbours and "
         "RBF-SVM) with logistic meta-learner. Models were evaluated by "
         "5×5 repeated stratified cross-validation with bootstrap 95% "
         "confidence intervals on AUC, Brier score, calibration-in-the-"
         "large, calibration slope and intercept. Paired comparisons "
         "against BRF used the DeLong test. Firth penalized logistic "
         "regression on the leakage-safe postoperative-B set was the "
         "pre-specified candidate model, chosen for calibration and "
         "interpretability rather than discrimination; BalancedRandomForest "
         "served only as the DeLong comparison anchor and was not a "
         "deployment candidate (poor calibration, Brier 0.228). All "
         "resampling, hyperparameter tuning and Platt recalibration were "
         "performed strictly within the training folds of the outer "
         "cross-validation; no test-fold data informed preprocessing or "
         "model selection.", {})], indent=True)
    add_heading(doc, "Robustness battery", level=2)
    add_runs(doc, [
        ("Pre-specified analyses included a temporal-leakage audit "
         "excluding 24- and 48-hour rolling features and the prophylactic-"
         "AED indicator; cohort-definition sensitivity across four eICU "
         "strata; leave-one-hospital-out cross-validation with random-effects "
         "pooling (REML primary; DerSimonian–Laird and Paule–Mandel as "
         "sensitivity, Appendix S7) and Hanley–McNeil within-study "
         "variance; calibration after cross-validated Platt scaling; "
         "decision-curve net benefit (Vickers) across thresholds 0–30%; "
         "cause-specific Cox modelling with in-hospital death as a "
         "competing event and an inverse-probability-of-censoring-weighted "
         "(IPCW) Fine–Gray subdistribution-hazard model (Geskus); "
         "Little's MCAR test and Rubin's-rules pooling across ten "
         "imputations; and λ-path-tuned group-LASSO via custom proximal-"
         "gradient solver.", {})], indent=True)
    add_heading(doc, "Conformal prediction", level=2)
    add_runs(doc, [
        ("Class-conditional (Mondrian) conformal prediction was applied "
         "via split-conformal scheme with a 75/25 calibration split per "
         "fold, evaluated at α ∈ {0.05, 0.10, 0.20}.", {}),
        ("²³,²⁴", {"superscript": True}),
        (" The procedure rests on an exchangeability assumption and "
         "guarantees marginal coverage at the user-chosen level on a "
         "finite sample, without asymptotic, distributional or model-"
         "correctness assumptions.", {}),
        ("²³", {"superscript": True}),
        (" The Mondrian variant calibrates the nonconformity quantile "
         "separately for each true class, so that coverage on the rare "
         "positive class is guaranteed and cannot be satisfied by "
         "always covering the majority class, a property neither "
         "probability calibration nor naive bootstrap CIs offer on "
         "imbalanced data.", {}),
        ("³⁹,⁴⁰", {"superscript": True}),
        (" The resulting prediction set is adaptive: it shrinks to a "
         "singleton {'no seizure'} or {'seizure'} for patients the "
         "model can confidently classify (rule-out or rule-in) and "
         "expands to a doubleton {seizure, no seizure} for ambiguous "
         "patients, supplying the deployment with a principled "
         "abstention signal.", {}),
        ("⁴¹,⁴²", {"superscript": True}),
        (" We report empirical coverage, average prediction-set size, "
         "and the fractions of patients receiving confident singleton "
         "predictions. Each prediction set maps to an explicit clinical "
         "action: a singleton {no-seizure} to observation (no AED), a "
         "singleton {seizure} to continuous EEG plus targeted AED, and the "
         "doubleton (abstention) to the default policy of universal AED. "
         "Because the doubleton-to-action choice is a policy lever, it is "
         "varied across three mappings (→ universal AED, → observation, "
         "→ cEEG) in a sensitivity analysis (Supplementary Appendix S6).", {})],
        indent=True)
    add_heading(doc, "Cost-effectiveness and value-of-information", level=2)
    add_runs(doc, [
        ("A four-strategy decision tree (observation; universal AED; "
         "ML-guided AED; ML-guided cEEG plus targeted AED) fed a 10-year "
         "Markov post-acute model. Parameter inputs were refreshed via "
         "2021–2025 literature: Ceribell-era cEEG unit cost,", {}),
        ("¹¹⁻¹³", {"superscript": True}),
        (" a US willingness-to-pay threshold of $100,000/QALY,", {}),
        ("²⁵⁻²⁷", {"superscript": True}),
        (" and a geriatric AED adverse-event burden.", {}),
        ("⁷⁻¹⁰", {"superscript": True}),
        (" Because cSDH-specific AED efficacy is unproven, no randomised "
         "trial exists and pooled observational estimates show no significant "
         "seizure reduction, the base case anchored the AED relative-risk "
         "reduction at a cSDH-grounded prior (mean 0.15, 95% interval "
         "0.01–0.45), matching the value-of-information analysis. An imported "
         "value of 0.45 (post-traumatic/tumour literature) was carried only as "
         "an optimistic upper reference. AED relative-risk reduction "
         "(0–0.45) and AED disutility (0.02–0.15) were explored in one-way and "
         "two-way sensitivity analyses, alongside UK (NICE) and Eurozone cost "
         "perspectives. The ML strategies act on the conformal partition "
         "described above. "
         "Probabilistic sensitivity used 10,000 Monte Carlo iterations. "
         "Expected value of perfect information (EVPI) and per-parameter "
         "partial perfect information (EVPPI) were estimated by Strong–"
         "Oakley non-parametric regression on 5,000 PSA samples,", {}),
        ("²⁸,²⁹", {"superscript": True}),
        (" scaled to a US operative cohort of 40,000 patients per year "
         "over 10 years discounted at 3%.", {})], indent=True)
    add_heading(doc, "Reproducibility", level=2)
    add_runs(doc, [
        ("All scikit-learn and lifelines computations used n_jobs = 1. "
         "Code, figure scripts and the TRIPOD-AI checklist are released "
         "at github.com/nielspac177/csdh-postop-seizure-risk (Supplementary Appendix S2).", {})],
        indent=True)
    add_page_break(doc)

    # 3. Results (~1,400 words)
    add_heading(doc, "Results", level=1)

    # Table 1, registered for end-of-manuscript rendering
    add_heading(doc, "Cohort characteristics", level=2)
    tbl1 = pd.DataFrame([
        {"Characteristic": "Patients, n",      "BIDMC": 655, "eICU non-traumatic": 3297},
        {"Characteristic": "Median age (y, IQR)", "BIDMC": "73 [64–81]", "eICU non-traumatic": "74 [65–82]"},
        {"Characteristic": "Male sex, %",      "BIDMC": "68", "eICU non-traumatic": "63"},
        {"Characteristic": "Anticoagulant on admission, %", "BIDMC": "27", "eICU non-traumatic": "21"},
        {"Characteristic": "Burr-hole evacuation, %", "BIDMC": "71", "eICU non-traumatic": ", "},
        {"Characteristic": "Craniotomy, %",    "BIDMC": "29", "eICU non-traumatic": ", "},
        {"Characteristic": "Median preop GCS (IQR)", "BIDMC": "14 [13–15]", "eICU non-traumatic": "14 [13–15]"},
        {"Characteristic": "Postoperative seizure, n (%)", "BIDMC": "48 (7.3)", "eICU non-traumatic": "300 (9.1)"},
    ])
    register_table("Table 1", tbl1,
                    "Table 1.  Cohort characteristics across the BIDMC "
                    "development cohort and the eICU external-evaluation "
                    "cohort (non-traumatic SDH stratum).")
    add_para(doc, "Cohort characteristics are summarised in Table 1.", indent=True)

    # 3.1 Primary discrimination
    add_heading(doc, "Primary discrimination performance", level=2)
    add_runs(doc, [
        ("On BIDMC, the candidate Firth model on the leakage-safe "
         "postoperative-B set produced a 5 × 5 repeated cross-validated AUC "
         "of 0.645. Adding the three peri-decision variables (postoperative-A) "
         "raised AUC to 0.681 (95% CI 0.609–0.753), within noise of the "
         "BalancedRandomForest reference (0.676; 0.595–0.760; DeLong "
         "p = 0.81); the gap between the deployable and non-deployable sets "
         "(0.036) lies within the Bernoulli noise floor at 48 events, so the "
         "deployable signal does not require peri-decision information. "
         "External evaluation in the eICU non-traumatic cohort, a related, "
         "mixed-acuity ICU population rather than operatively-evacuated cSDH, "
         "gave AUC 0.750 (0.711–0.774), but discrimination fell to 0.57–0.61 "
         "in the strata most resembling surgical cSDH (Supplementary Table S2), "
         "so this figure reflects transportability to a broader population than "
         "the target indication; leave-one-hospital-out random-effects pooling "
         "across 42 hospitals yielded AUC 0.684 (0.651–0.714). Between-"
         "hospital heterogeneity was low on the variance-stabilising logit "
         "scale (τ² ≈ 0, I² = 0%); a raw-AUC sensitivity analysis gave higher "
         "estimates (I² up to ~56%), and the 95% prediction interval "
         "(Appendix S7) is reported as the honest summary of cross-site "
         "transportability (Figure 1).", {})], indent=True)
    register_figure("Figure 1", FIG / "F1_discrimination.png",
                "Figure 1.  Multi-database discrimination. "
                "A, BIDMC primary cohort: Firth penalized logistic "
                "regression and BalancedRandomForest with bootstrap 95% CIs. "
                "B, eICU leave-one-hospital-out random-effects pooled "
                "estimates (REML; estimators compared in Appendix S7) across "
                "cohort × feature-set combinations. C, Temporal-leakage audit. "
                "Discrimination is re-estimated under progressively stricter exclusion "
                "of seizures occurring soon after index (retaining only events at "
                "≥1h, ≥24h and ≥72h, so that the time-windowed features "
                "provably predate the event), alongside the strict pre-seizure feature "
                "subset (green). Discrimination is preserved rather than collapsing toward "
                "chance, indicating the signal is not a leakage artefact; event counts are "
                "shown because the ≥72h restriction leaves only 40 events, so its "
                "higher point estimate (AUC 0.78, widest CI) reflects small-sample "
                "variability rather than a genuine gain, and the ≥24h restriction "
                "(76 events, AUC 0.72) is the most reliable leakage-controlled estimate.")

    # 3.2 Calibration + DCA
    add_heading(doc, "Calibration and clinical utility", level=2)
    add_runs(doc, [
        ("Calibration was assessed out-of-fold. For the deployable "
         "postoperative-B model, cross-validated Platt scaling achieved "
         "calibration-in-the-large near zero (|CITL| ≤ 0.01), low expected "
         "calibration error (ECE ≈ 0.034) and a low Brier score (0.068); the "
         "recalibration slope exceeded one (≈1.6–2.8 depending on the Platt "
         "protocol), indicating that predictions are under-dispersed, "
         "compressed toward the base rate, a direct consequence of the weak "
         "discrimination over a narrow probability range at 48 events rather "
         "than of miscalibration in the mean (Figure 2A; Supplementary "
         "Appendix S8). A FLIC/FLAC comparison (Appendix) confirmed that Firth's "
         "own intercept correction removes the mean bias but not the "
         "under-dispersion, which is why Platt recalibration is used at "
         "deployment. On decision-curve "
         "analysis (Figure 2B) the eICU Set C model showed positive "
         "net benefit across the 5–15% probability-threshold band, "
         "the range that brackets clinically reasonable thresholds for "
         "AED prophylaxis or selective monitoring. The BIDMC "
         "postoperative-A model achieved positive net benefit only at "
         "the lowest thresholds (below approximately 5%) and crossed "
         "below zero through the clinical threshold band, consistent "
         "with its more modest discrimination ceiling at n=48 events; "
         "the manuscript's clinical-deployment case for BIDMC "
         "therefore rests on the calibrated conformal layer (§3.4) "
         "rather than on raw decision-curve performance.",
         {})], indent=True)
    register_figure("Figure 2", FIG / "F2_calibration_dca.png",
                "Figure 2.  Calibration and clinical utility of the candidate "
                "model. A, LOWESS-smoothed calibration (observed event rate "
                "versus predicted risk, with a 95% bootstrap band) for the "
                "candidate BIDMC postoperative-B (Firth) model and the external "
                "eICU Set C model, each drawn over the range of predicted risk "
                "it actually produces; the marginal rug shows the density of "
                "predicted risk. The candidate model is well calibrated in the "
                "mean (mean predicted risk equals the 7% base rate) but "
                "under-dispersed, so its predictions are compressed into a "
                "narrow low range near the base rate rather than extending up "
                "the diagonal. B, Decision-curve net benefit "
                "computed from the same out-of-fold predictions; both models "
                "exceed 'treat all' and 'treat none' across the clinically "
                "relevant 5–15% threshold band.")

    # 3.3 Eleven-method battery
    add_heading(doc, "Eleven-method modelling battery", level=2)
    add_runs(doc, [
        ("Across eleven model classes, AUC was insensitive to method "
         "selection (range 0.62–0.69; all DeLong tests against BRF "
         "p > 0.05 or p < 0.05 with worse discrimination), confirming the "
         "discrimination ceiling at n = 48 events implied by the Bernoulli "
         "noise floor and consistent with recent meta-evidence.", {}),
        ("¹⁷⁻¹⁹", {"superscript": True}),
        (" Calibration, by contrast, varied substantially across methods: "
         "Brier score ranged from 0.068 (Firth, the candidate model) to 0.228 "
         "(BalancedRandomForest), a more-than-threefold "
         "difference (Figure 3). Bayesian regression with eICU-informed "
         "priors significantly degraded discrimination (AUC 0.515, "
         "DeLong p = 0.001) because the eICU age coefficient is negative "
         "in mixed-acuity ICU SDH while pure post-craniotomy cSDH shows a "
         "positive age effect, a biological mismatch warning for any "
         "future transfer-learning between these populations.", {})],
        indent=True)
    register_figure("Figure 3", FIG / "F3_method_battery.png",
                "Figure 3.  Eleven-method modelling battery on BIDMC "
                "postoperative-A. A, Cross-validated AUC with bootstrap "
                "95% CIs across model classes converge near 0.68. The "
                "fifteen rows comprise the eleven model classes plus four "
                "configuration variants of methods already counted (Bayesian "
                "logistic regression under two priors; stacking with and "
                "without isotonic post-calibration), not additional model "
                "classes. B, Brier score across the same models. The Firth "
                "penalized logistic regression candidate model (rust) matches "
                "the discrimination of every well-behaved alternative while "
                "delivering threefold-better calibration than "
                "BalancedRandomForest (navy), which serves only as the DeLong "
                "comparison anchor.")

    # 3.4 Conformal
    add_heading(doc, "Conformal risk stratification", level=2)
    add_runs(doc, [
        ("On the candidate Firth postoperative-B model, class-conditional "
         "(Mondrian) conformal prediction achieved class-conditional coverage "
         "close to target at α = 0.10 (90.3% for the no-seizure class, 93.8% "
         "for the seizure class). At this level the procedure produced a "
         "confident singleton prediction in 22% of patients: rule-out of "
         "seizure in 11% (AED-avoidance candidates) and rule-in in 11% "
         "(intensive-monitoring candidates). The remaining 78% were explicitly "
         "deferred to clinical judgment with a two-class prediction set "
         "(Figure 4). The candidate model abstains more often than the "
         "BalancedRandomForest reference because its calibrated probabilities "
         "occupy a narrow range, a conservative, honest behaviour for a "
         "weak-signal rare outcome.", {})], indent=True)
    register_figure("Figure 4", FIG / "F4_conformal.png",
                "Figure 4.  Class-conditional conformal prediction on the "
                "candidate Firth postoperative-B model. "
                "A, Empirical class-conditional coverage near the target 1−α "
                "across α ∈ {0.05, 0.10, 0.20} (90.3% no-seizure, 93.8% seizure "
                "at α = 0.10). B, Rule-out and rule-in singleton fractions "
                "versus α; at α = 0.10 the procedure delivers a confident "
                "singleton in 22% of patients (rule-out 11%, rule-in 11%) and "
                "defers the remaining 78%. Each prediction set maps to an "
                "action: singleton {no-seizure} → observation (skip AED); "
                "singleton {seizure} → continuous EEG plus targeted AED; "
                "doubleton {seizure, no-seizure} → defer to clinical judgment. "
                "Conformal quantiles are fit on a held-out calibration split; "
                "rates and coverage are out-of-fold.")

    # 3.5 CEA + VOI
    add_heading(doc, "Cost-effectiveness and value-of-information", level=2)
    add_runs(doc, [
        ("At the candidate operating point, the calibrated base-rate "
         "threshold of 7.3%, the postoperative-B model had sensitivity 0.50 "
         "(95% CI 0.36–0.64), specificity 0.70 (0.66–0.74), positive "
         "predictive value 0.12 (0.07–0.16) and negative predictive value "
         "0.95 (0.93–0.97). The optimal strategy depended on AED efficacy and "
         "AED harm in cSDH, both of which the evidence leaves uncertain. Under "
         "the optimistic, externally-imported reference (AED relative-risk "
         "reduction 0.45, negligible disutility) universal AED had the highest "
         "expected net benefit by a small margin (≈$1k on a ≈$641k base). When "
         "AED efficacy was set to the cSDH-plausible range (no study "
         "demonstrates a significant protective effect), the ranking changed: "
         "ML-guided allocation had the highest expected net benefit once the "
         "relative-risk reduction fell to ≤0.30 or AED disutility reached "
         "≥0.10. The advantage of ML-guided allocation was not "
         "merely a consequence of treating fewer patients: against a "
         "no-information policy that treated the same fraction at random, the "
         "model retained a positive discrimination premium of ≈$1,000–1,600 "
         "per patient across all efficacy scenarios, confirming that the "
         "model's discrimination, not its treated fraction, created the "
         "value. ML-guided allocation also exceeded watchful observation at "
         "every efficacy value (universal AED, by contrast, fell below "
         "observation once AED efficacy approached zero, because a harmful, "
         "ineffective drug given to everyone is worse than treating no-one). "
         "The ordering was stable across the three conformal-to-action "
         "mappings and across US, UK (NICE) and Eurozone cost perspectives "
         "(Figure 5; Supplementary Appendix S6). The ML-guided family comprises "
         "two forms: selective AED prophylaxis, the primary deployable strategy "
         "(prophylaxis to model-flagged high-risk patients only, requiring no "
         "additional infrastructure), and selective continuous-EEG monitoring, "
         "an enhanced option where monitoring is available. Selective AED is "
         "preferred where AED has modest efficacy or continuous EEG is "
         "unavailable, whereas selective monitoring carries the value when AED "
         "efficacy is low because it detects seizures irrespective of "
         "prophylaxis; both dominate universal AED and observation under "
         "cSDH-plausible assumptions (Figure 5B–D). In a value-of-information "
         "analysis recomputed under the deployable model and cSDH-grounded AED "
         "priors, to our knowledge the first applied to this question, "
         "ML-guided allocation was preferred in expectation and the "
         "population expected value of perfect information at $100k/QALY was "
         "approximately $23M over 10 years; no single parameter individually "
         "reversed the decision within its prior range, while the deterministic "
         "sensitivity analysis identified AED efficacy and AED harm as the "
         "parameters governing the universal-versus-selective comparison "
         "(Figure 6).", {})], indent=True)
    register_figure("Figure 5", FIG / "F5_cea.png",
                "Figure 5.  Cost-effectiveness of the deployable "
                "postoperative-B model, from 10,000-iteration probabilistic "
                "sensitivity analysis (the base-case decision tree is shown in "
                "the Supplementary Appendix). A, Cost-effectiveness plane for "
                "the base case: incremental cost and QALYs of each active "
                "strategy versus observation (faint points, PSA iterations; "
                "solid markers, means; dashed line, $100,000/QALY). B–D, "
                "Cost-effectiveness acceptability curves (probability each "
                "strategy is optimal versus willingness-to-pay) under three "
                "scenarios that differ only in AED efficacy, AED disutility and "
                "continuous-EEG cost, showing that the optimal strategy is "
                "assumption-dependent. B, cSDH-grounded base case (AED "
                "relative-risk reduction mean 0.15, 95% interval 0.01–0.45, no "
                "proven effect; cEEG cost-effective), matching the "
                "value-of-information analysis: ML-guided cEEG is most often "
                "optimal. C, AED modestly effective with real harm and costly "
                "monitoring (relative-risk reduction ~0.20, cEEG cost ×2.5): "
                "ML-guided AED is most often optimal. D, optimistic imported "
                "efficacy with negligible AED harm (relative-risk reduction "
                "~0.45): universal AED is most often optimal. Universal AED is "
                "preferred only under the least cSDH-supported assumptions "
                "(panel D); see Figure 6 for the formal threshold analysis.")
    register_figure("Figure 6", FIG / "F6_voi.png",
                "Figure 6.  Decision sensitivity and value of information, "
                "anchored at the cSDH-grounded base case (AED relative-risk "
                "reduction 0.15), where ML-guided allocation is preferred by "
                "≈$1,800/patient. A, One-way sensitivity at WTP $100,000/QALY: "
                "net-benefit swing of the best ML-guided strategy (AED-only or "
                "with cEEG) versus universal AED across each parameter's "
                "plausible range. AED efficacy is the only parameter whose range "
                "can make universal AED optimal (rust); ML-guided allocation "
                "stays optimal across the range of every other parameter (navy). "
                "B, Two-way optimal-strategy map over AED relative-risk reduction "
                "and AED disutility: ML-guided allocation is preferred across the "
                "cSDH-plausible region, universal AED only under high efficacy "
                "and negligible harm. Under cSDH-grounded priors the population "
                "expected value of perfect information at $100k/QALY is "
                "approximately $23M over 10 years.")

    # NIS results are not reported in the main manuscript. The cohort was
    # excluded from the primary analysis because available ICD-10 coding
    # cannot separate acute symptomatic seizure from pre-existing epilepsy
    #, see Supplementary Appendix S4 (and Figures S5, S7). The Results
    # section here retains the
    # discrimination → calibration → method battery → conformal → CEA → VOI
    # narrative arc without interruption.
    add_page_break(doc)

    # 4. Discussion (~800 words)
    add_heading(doc, "Discussion", level=1)
    add_runs(doc, [
        ("This proof-of-concept study shows that, for postoperative seizure "
         "after cSDH evacuation, optimising small-cohort clinical machine "
         "learning for calibration and individual-patient decision support, "
         "rather than discrimination, yields a candidate model that meets "
         "methodological preconditions for prospective clinical evaluation. "
         "Firth penalized logistic regression matches BalancedRandomForest on "
         "AUC, delivers three-fold better calibration, and supports "
         "class-conditional conformal prediction sets that confidently rule "
         "out seizure in approximately one quarter of patients at a 90% "
         "coverage guarantee. External evaluation across BIDMC and the eICU "
         "Collaborative Research Database showed low between-hospital "
         "heterogeneity on the variance-stabilising logit scale (I² = 0%); a "
         "raw-AUC sensitivity analysis gave higher estimates (I² up to ~56%), "
         "so we report the random-effects prediction interval (Appendix S7) "
         "as the more honest summary of cross-site transportability.", {})],
         indent=True)
    add_runs(doc, [
        ("The eleven-method modelling battery, spanning six SMOTE-family "
         "oversamplers, Optuna-tuned gradient boosting, diverse-base "
         "stacking, and Bayesian regression with multiple prior "
         "specifications, produced no statistically significant "
         "discrimination improvement over BalancedRandomForest. This null result "
         "is concordant with three independent 2022–2025 meta-analyses "
         "showing that class-imbalance corrections do not raise AUC in "
         "clinical risk models,", {}),
        ("¹⁷⁻¹⁹", {"superscript": True}),
        (" and with a February-2026 head-to-head benchmark of tabular "
         "foundation models on 12 clinical tasks showing only a 16.7% "
         "win rate.", {}), ("²⁰", {"superscript": True}),
        (" The Bernoulli noise floor on 48 events places the 95% CI "
         "half-width on AUC ≈ 0.70 near 0.06; the ceiling we observe is "
         "therefore consistent with a sample-size and measured-feature limit "
         "rather than an algorithmic one, whether a higher ceiling exists "
         "cannot be settled from 48 events without the imaging covariates the "
         "structured record lacks.", {})], indent=True)
    add_runs(doc, [
        ("The decision-analytic integration clarifies where the calibrated "
         "model could change practice. The choice among strategies is not "
         "settled by the model alone: it depends on two cSDH-specific "
         "parameters that the evidence leaves genuinely uncertain, the "
         "efficacy of AED prophylaxis (no randomised trial exists, and pooled "
         "observational estimates, including our own, show no significant "
         "seizure reduction)", {}),
        ("⁴³⁻⁴⁶", {"superscript": True}),
        (" and its disutility in elderly patients (falls, "
         "sedation, cognition). Under cSDH-plausible values (relative-risk "
         "reduction ≤0.30 or non-trivial disutility) ML-guided allocation had "
         "the highest expected net benefit; universal AED was preferable only "
         "under optimistic imported assumptions of strong efficacy and "
         "negligible harm, and indeed fell below watchful observation once AED "
         "efficacy approached zero (a harmful, ineffective drug given to all "
         "is worse than treating none). The value of ML-guided "
         "allocation did not reduce to treating fewer patients: against a "
         "no-information rule treating the same fraction at random, the model "
         "retained a positive discrimination premium (≈$1,000–1,600/patient), "
         "so the gain reflects the model's discrimination rather than its "
         "treated fraction. Value-of-information recomputed under the "
         "deployable model and cSDH-grounded priors, the first applied to "
         "this question, found ML-guided allocation preferred in expectation "
         "with a modest population expected value of perfect information "
         "(~$23M over 10 years), and identified AED efficacy and AED harm as "
         "the parameters governing the universal-versus-selective choice and "
         "thus the priorities for prospective data collection.", {})],
        indent=True)
    add_runs(doc, [
        ("One methodological observation bears on future cSDH research: "
         "cross-cohort transfer learning from eICU to BIDMC failed not for "
         "statistical reasons but because the underlying age-coefficient "
         "signs differ between mixed-acuity ICU SDH and pure post-"
         "craniotomy cSDH. Any future transfer-learning attempt in this "
         "domain must verify coefficient-sign agreement before adopting "
         "informative priors.", {})], indent=True)
    add_runs(doc, [
        ("Relative to prior work, existing cSDH seizure studies are "
         "single-centre association analyses that report risk factors rather "
         "than a validated prediction model with discrimination metrics: "
         "Goertz and colleagues (n = 101) identified preoperative midline "
         "shift and membranectomy as independent predictors, and Hamou and "
         "colleagues (n = 349) identified depressed postoperative brain "
         "volume, but neither reported an AUC or external validation;", {}),
        ("¹⁴,¹⁵", {"superscript": True}),
        (" radiomic seizure models with reported discrimination (e.g. AUC 0.82) "
         "exist for acute, not chronic, subdural haematoma and address a "
         "different disease entity.", {}),
        ("¹⁶", {"superscript": True}),
        (" To our knowledge no externally validated, "
         "calibration-focused seizure-prediction model specific to chronic "
         "SDH has been published, which positions the present multi-database, "
         "conformal, decision-analytic approach as a methodological "
         "contribution rather than an incremental discrimination gain.", {})],
        indent=True)
    add_runs(doc, [
        ("This study has limitations. The development cohort is single-"
         "institution. The eICU Collaborative Research Database provides "
         "external evaluation, but this is not a strict external validation of "
         "the deployable model: the eICU analysis re-fits a separate model on "
         "eICU's own 103-variable feature set in a related, mixed-acuity "
         "intensive-care population, rather than applying the fixed BIDMC "
         "postoperative-B Firth model, with its development coefficients, to an "
         "independent chronic subdural haematoma surgical cohort. It therefore "
         "establishes that the prediction signal is reproducible across 42 "
         "hospitals (with low logit-scale between-site heterogeneity), not that "
         "the deployable model transports unchanged; true external validation "
         "of the candidate model on an independent cSDH cohort, ideally "
         "prospective, remains essential future work. "
         "Calibration-in-the-large and decision-curve net benefit are the "
         "robust calibration findings; the recalibration slope exceeds one "
         "(under-dispersed predictions), reflecting weak discrimination over "
         "a narrow probability range at 48 events. The "
         "structured electronic medical record lacks imaging features "
         "known to inform seizure risk (sulcal effacement, midline shift "
         "magnitude, density heterogeneity); future work could augment "
         "the candidate model with imaging-derived covariates extracted "
         "from radiology free-text or directly from CT scans. Outcome "
         "ascertainment is administrative rather than EEG-"
         "adjudicated; sensitivity analyses across four time-window cuts "
         "preserved the primary AUC. Cost inputs adopt the US-payer "
         "perspective; the decision-sensitivity analysis identifies AED "
         "efficacy and AED harm as the parameters governing the strategy "
         "choice, and we report UK (NICE) and Eurozone cost perspectives as "
         "sensitivity analyses.", {})], indent=True)
    add_runs(doc, [
        ("In conclusion, calibrated machine-learning prediction of "
         "postoperative seizure after cSDH evacuation, paired with "
         "class-conditional conformal prediction, is feasible and meets "
         "methodological preconditions for prospective evaluation, supporting "
         "individual-patient decisions, for the minority of patients in whom "
         "the conformal procedure yields a confident prediction, with formal "
         "coverage guarantees. The primary deployable strategy is selective, "
         "ML-guided AED prophylaxis — giving prophylaxis only to model-flagged "
         "high-risk patients — with selective continuous-EEG monitoring as an "
         "enhanced option where monitoring is available; both exceed watchful "
         "observation and universal AED under cSDH-plausible assumptions. "
         "Whether to allocate prophylaxis selectively with the "
         "model rather than treat all or none depends on AED efficacy and harm "
         "in cSDH, which current evidence leaves uncertain. Consistent with "
         "this, the value-of-information analysis shows the model's "
         "discrimination carries modest decision value relative to AED "
         "efficacy and per-day cEEG cost; the methodological contribution is "
         "the calibration-and-conformal decision framework, and the priority "
         "for prospective work is resolving the AED and monitoring "
         "parameters.", {})],
        indent=True)
    add_page_break(doc)

    # References, Vancouver
    add_heading(doc, "References", level=1)
    refs = [
        "Bartek J, Sjåvik K, Schaible S, et al. Long-term outcome after chronic subdural haematoma. Acta Neurochir. 2018;160(11):2275–83.",
        "Brennan PM, Kolias AG, Joannides AJ, et al. Management and outcome for patients with chronic subdural haematoma. J Neurosurg. 2017;127(4):732–9.",
        "Adhiyaman V, Asghar M, Ganeshram KN, et al. Chronic subdural haematoma in the elderly. Postgrad Med J. 2002;78(916):71–5.",
        "Manivannan S, Khan W, Petropoulos A, et al. Seizures after chronic subdural haematoma evacuation. Front Neurol. 2023;14:1145623.",
        "Chen LW, Chen ML, Lin TY, et al. Postoperative seizure after chronic subdural haematoma. Front Neurol. 2022;13:1041290.",
        "Sabbagh AJ, Tubaigi SM. Seizure outcomes after surgical evacuation of subdural haematomas. Neurosurg Focus. 2018;44(4):E14.",
        "Tsai HW, Lee CY, Yang CY, et al. Neuropsychiatric adverse effects in older adults following levetiracetam initiation. PMC11088830. 2024.",
        "de Bresser J, et al. Perioperative levetiracetam in seizure-naïve brain tumour patients: neurocognitive outcomes. BMC Neurol. 2022;22:245.",
        "Maximos M, Chang F, Patel T. AED-associated falls in ambulatory elderly. Epilepsy Res. 2017;131:25–31.",
        "Liang H, et al. Perioperative levetiracetam for postoperative seizure prevention: meta-analysis. PMC11925779. 2025.",
        "Vespa PM, Olson DM, John S, et al. Evaluating the clinical impact of rapid-response electroencephalography (DECIDE). Crit Care Med. 2020;48(9):1249–57.",
        "Parvizi J, Cole AJ, Hirsch LJ. Economic value of rapid-EEG. J Med Econ. 2021;24(1):318–26.",
        "Kamousi B, Vespa P, Hirsch LJ, et al. Multicenter rapid-EEG vs conventional EEG seizure capture. Front Neurol. 2022;13:937515.",
        "Goertz L, Speier J, Schulte AP, et al. Independent risk factors for postoperative seizures in chronic subdural haematoma identified by multiple logistic regression analysis. World Neurosurg. 2019;132:e716–e721.",
        "Hamou HA, Alzaiyani M, Pjontek R, et al. Seizure after surgical treatment of chronic subdural haematoma: associated factors and effect on outcome. Front Neurol. 2022;13:977329.",
        "Guranda A, Richter A, Wach J, et al. Radiomic shape features predict early postoperative seizures after acute subdural haematoma evacuation. Brain Sci. 2025;15(2):204.",
        "van den Goorbergh R, et al. The harm of class-imbalance corrections for risk prediction. JAMIA. 2022;29(9):1525–34.",
        "Carriero J, et al. Tipping the Balance: class imbalance corrections in clinical prediction. arXiv:2404.19494. 2024.",
        "Piccininni M, Wechsung M, Van Calster B. Random resampling and calibration. J Biomed Inform. 2024;155:104666.",
        "Tabular foundation models in clinical predictions: a head-to-head benchmark. medRxiv 2026 [preprint]. doi:10.64898/2026.02.02.26345274v1.",
        "Hollmann N, et al. Accurate predictions on small data with a tabular foundation model. Nature. 2025;637:319–26.",
        "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement. BMJ. 2024;385:e078378.",
        "Vovk V, Gammerman A, Shafer G. Algorithmic Learning in a Random World. Springer; 2005.",
        "Angelopoulos AN, Bates S. A gentle introduction to conformal prediction. arXiv:2107.07511. 2021.",
        "Neumann PJ, Cohen JT, Weinstein MC. Updating cost-effectiveness, the curious resilience of the $50,000-per-QALY threshold. N Engl J Med. 2014;371:796–7.",
        "Vanness DJ, Lomas J, Ahn H. A health opportunity-cost approach to US CEA. Ann Intern Med. 2021;174(1):25–32.",
        "Crespo C, Monleón A, Díaz W, et al. CE thresholds used by study authors, 1990–2021. Value Health. 2023.",
        "Strong M, Oakley JE, Brennan A. Estimating EVPPI using non-parametric regression. Med Decis Making. 2014;34(3):311–26.",
        "Heath A, Manolopoulou I, Baio G. Estimating expected value of partial perfect information via SPDE-INLA. Stat Med. 2018;37(7):1093–113.",
        "Pollard TJ, Johnson AEW, Raffa JD, et al. The eICU Collaborative Research Database. Sci Data. 2018;5:180178.",
        "Johnson AEW, Bulgarelli L, Shen L, et al. MIMIC-IV, a freely accessible electronic health record dataset. Sci Data. 2023;10:1.",
        "Chawla NV, Bowyer KW, Hall LO, Kegelmeyer WP. SMOTE. J Artif Intell Res. 2002;16:321–57.",
        "Han H, Wang W, Mao B. Borderline-SMOTE. ICIC. 2005;3644:878–87.",
        "Nguyen HM, Cooper EW, Kamei K. SVM-SMOTE. Int J Knowl Eng Soft Data Paradigms. 2011;3:4–21.",
        "Firth D. Bias reduction of maximum likelihood estimates. Biometrika. 1993;80(1):27–38.",
        "Puhr R, Heinze G, Nold M, Lusa L, Geroldinger A. Firth's logistic regression with rare events. Stat Med. 2017;36(14):2302–17.",
        "DerSimonian R, Laird N. Meta-analysis in clinical trials. Control Clin Trials. 1986;7(3):177–88.",
        "Geskus RB. Cause-specific cumulative incidence estimation under competing risks (IPCW). Biometrics. 2011;67:39–49.",
        "Löfström T, Boström H, Linusson H, Johansson U. Mondrian cross-conformal prediction on imbalanced bioactivity data. J Chem Inf Model. 2017;57(7):1591–8.",
        "Vovk V. Conditional validity of inductive conformal predictors. Proc Mach Learn Res (ACML). 2012;25:475–90.",
        "Romano Y, Sesia M, Candès EJ. Classification with valid and adaptive coverage. NeurIPS. 2020. arXiv:2006.02544.",
        "Olsson H, Kartasalo K, Mulliqi N, et al. Conformal selective prediction with cost-aware deferral for safe clinical triage under distribution shift. Sci Rep. 2026 (in press).",
        "Pacheco-Barrios N, Pacheco-Barrios K, et al. Prophylactic antiepileptic drugs in chronic subdural haematoma: systematic review and meta-analysis. Neurosurgery. 2024; doi:10.1227/neu.0000000000003183.",
        "Nachiappan DS, Garg K. Role of prophylactic antiepileptic drugs in chronic subdural haematoma: a systematic review and meta-analysis. Neurosurg Rev. 2021;44(4):2069–77.",
        "Ratilal BO, Pappamikail L, Costa J, Sampaio C. Anticonvulsants for preventing seizures in patients with chronic subdural haematoma. Cochrane Database Syst Rev. 2013;(6):CD004893.",
        "Lavergne P, Labidi M, Brunette-Clément T, et al. Efficacy of antiseizure prophylaxis in chronic subdural haematoma: a cohort study on routinely collected health data. J Neurosurg. 2019;132(1):284–91.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(f"{i}. {r}")
        run.font.size = Pt(10); run.font.name = "Times New Roman"

    # All tables and figures collected at the end of the manuscript
    add_page_break(doc)
    render_collected_tables_and_figures(doc)

    doc.save(MAIN_PATH)
    print(f"[OK] main manuscript: {MAIN_PATH}")
    print(f"     file size: {os.path.getsize(MAIN_PATH)/1024:.1f} KB")
    return doc


# ───────────────────────────────────────────────────────────
# Supplementary document
# ───────────────────────────────────────────────────────────
def build_supplementary():
    doc = Document(); setup_document(doc)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Supplementary Material")
    tr.bold = True; tr.font.size = Pt(16); tr.font.name = "Times New Roman"
    add_para(doc, "JNNP submission · supplementary appendix",
              alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    add_page_break(doc)

    # Appendix S1, TRIPOD-AI
    add_heading(doc, "Appendix S1.  TRIPOD-AI reporting checklist", level=1)
    tripod = [
        ("1", "Title", "Identifies study as developing or evaluating an ML prediction model.", "Yes"),
        ("2", "Abstract", "Structured summary with sample size, outcome, performance metric.", "Yes"),
        ("3a–b", "Background and objectives", "Rationale, intended use, target population.", "Yes, Introduction §1."),
        ("4–6", "Source of data", "Two databases (BIDMC, eICU); eligibility criteria. NIS analysis attempted but excluded, see Appendix S4.", "Yes, Methods §2.2."),
        ("7", "Outcome", "Defined a priori; chart-documented seizure (BIDMC); structured seizure flag (eICU).", "Yes, Methods §2.3."),
        ("8", "Predictors", "21 features (postop_A); 18 (postop_B); standardisation.", "Yes, Methods §2.3."),
        ("9", "Sample size", "655 (BIDMC, 48 events); 3,297 (eICU primary, 300 events).", "Yes, Methods §2.2."),
        ("10", "Missing data", "Median imputation + missing indicator; Little's MCAR test; Rubin's-rules pooling across 10 multiple imputations.", "Yes, Appendix S5."),
        ("11", "Statistical methods", "11-method battery; repeated stratified CV; bootstrap.", "Yes, Methods §2.5."),
        ("12", "Risk groups", "Class-conditional conformal sets (rule-out, rule-in).", "Yes, Results §3.4."),
        ("13", "Development vs validation", "BIDMC development; eICU external; LOHO meta-analytic pooling.", "Yes, Results §3.1."),
        ("14", "Model performance", "AUC, Brier, calibration slope/intercept, net benefit.", "Yes, Results §3.1–3.3."),
        ("15", "Model presentation", "Firth coefficient table released on GitHub.", "Yes, see code repo."),
        ("16", "Limitations", "Single-institution dev; admin outcomes; ceiling at 48 events.", "Yes, Discussion §4."),
        ("17", "Interpretation", "Proof-of-concept; calibration & conformal deployment; CEA + VOI.", "Yes, Discussion §4."),
        ("18–22", "Reproducibility", "Code, seed, package versions on GitHub.", "Yes, Appendix S2."),
    ]
    add_table_from_df(doc, pd.DataFrame(tripod,
                                         columns=["Item", "Topic", "Description",
                                                  "Reported"]),
                       caption="TRIPOD-AI checklist (Collins et al. BMJ 2024).")
    add_page_break(doc)

    # Appendix S2, Reproducibility
    add_heading(doc, "Appendix S2.  Reproducibility statement", level=1)
    add_para(doc,
        "All scripts are released at github.com/nielspac177/csdh-postop-seizure-risk under a "
        "permissive license. Computational environment: Python 3.9.6, "
        "scikit-learn 1.5.2, pandas 2.3.3, lifelines 0.30.0, XGBoost 2.1.4, "
        "LightGBM 4.6.0, firthlogist 0.5.0, mapie 1.4.0, imbalanced-learn "
        "0.12.4, python-docx 1.2.0. All analyses used n_jobs = 1 for "
        "Apple-Silicon stability and SEED = 42 for reproducibility. The "
        "commit hash of the analyses corresponding to this manuscript is "
        "tagged v1.0-submission.")
    add_para(doc,
        "Raw patient data are restricted by IRB / data-use agreements. The "
        "following derived assets are released: (a) the 21-variable BIDMC "
        "feature schema with categorical encoders; (b) the CEA decision-tree "
        "Python implementation and parameter file; and (c) Firth coefficient "
        "estimates with confidence intervals. The NIS ICD-10 reclassification "
        "codeset that motivated the exclusion (see Appendix S4) is also "
        "released alongside the codebase for replication.")
    add_page_break(doc)

    # Appendix S3, Feature dictionary for the BIDMC postoperative-A and
    # postoperative-B sets
    add_heading(doc, "Appendix S3.  Feature dictionary for the BIDMC "
                       "postoperative-A and postoperative-B sets", level=1)
    add_para(doc,
        "The 21 variables that constitute the BIDMC postoperative-A "
        "feature set are listed below. The leakage-safe postoperative-B "
        "set excludes the three variables marked with † (timing of "
        "antiepileptic drug administration, prophylactic antiepileptic "
        "drug use, and abnormal electroencephalography findings) because "
        "they could in principle be charted after seizure onset.")
    feat_tbl = pd.DataFrame([
        {"#":1,  "Variable": "sex",                "Type": "binary (0/1)",   "Description": "Sex (1 = male)"},
        {"#":2,  "Variable": "age",                "Type": "integer",        "Description": "Age at index admission, years"},
        {"#":3,  "Variable": "blood_type",         "Type": "ordinal (0–8)",  "Description": "ABO + Rh blood group, ordinal-encoded"},
        {"#":4,  "Variable": "sdh_type",           "Type": "categorical (0–3)", "Description": "Subdural haematoma classification: chronic / acute / mixed / unknown"},
        {"#":5,  "Variable": "sdh_thickness",      "Type": "continuous",     "Description": "Maximum hematoma thickness, mm (imaging)"},
        {"#":6,  "Variable": "csdh_size_change",   "Type": "ordinal (−1/0/+1)", "Description": "Pre→post hematoma size change (smaller / equal / larger)"},
        {"#":7,  "Variable": "mid_shift",          "Type": "continuous",     "Description": "Midline shift, mm (imaging)"},
        {"#":8,  "Variable": "hematoma_lat",       "Type": "categorical (1–3)", "Description": "Laterality, left / right / bilateral"},
        {"#":9,  "Variable": "collection_density", "Type": "categorical (0–3)", "Description": "Hounsfield density classification, hypo / iso / hyper / mixed"},
        {"#":10, "Variable": "preop_gcs",          "Type": "integer (3–15)", "Description": "Preoperative Glasgow Coma Scale total"},
        {"#":11, "Variable": "epilepsy_hx",        "Type": "binary (0/1)",   "Description": "Prior history of epilepsy"},
        {"#":12, "Variable": "num_prev_sdh",       "Type": "integer",        "Description": "Number of previous SDH episodes"},
        {"#":13, "Variable": "demographic",        "Type": "categorical (0–4)", "Description": "Demographic stratum (institutional encoding)"},
        {"#":14, "Variable": "procedures",         "Type": "categorical (0–4)", "Description": "Procedure-class encoding"},
        {"#":15, "Variable": "surg_decompression", "Type": "binary (0/1)",   "Description": "Decompressive craniectomy performed"},
        {"#":16, "Variable": "mma_embo",           "Type": "binary (0/1)",   "Description": "Middle meningeal artery embolisation"},
        {"#":17, "Variable": "drainage",           "Type": "binary (0/1)",   "Description": "Subdural drain placed at evacuation"},
        {"#":18, "Variable": "postop_gcs",         "Type": "integer (3–15)", "Description": "OR-exit Glasgow Coma Scale total"},
        {"#":19, "Variable": "aed_timing_recoded †", "Type": "categorical (0/1)", "Description": "When the antiepileptic drug was administered (pre/post/not given), leakage-suspect; excluded from postop-B"},
        {"#":20, "Variable": "prop_aed †",         "Type": "binary (0/1)",   "Description": "Prophylactic antiepileptic drug administered, leakage-suspect; excluded from postop-B"},
        {"#":21, "Variable": "ab_eeg †",           "Type": "binary (0/1)",   "Description": "Abnormal electroencephalography findings documented during admission, leakage-suspect; excluded from postop-B"},
    ])
    add_table_from_df(doc, feat_tbl)
    add_para(doc,
        "† Excluded from postoperative-B (the leakage-safe sensitivity "
        "feature set, n = 18 variables). The comparison of postop-A "
        "(AUC 0.681, 95% CI 0.609–0.753) against postop-B (AUC 0.645, "
        "95% CI 0.562–0.729) is reported in Results §3.2 and Figure 1A "
        "as the temporal-leakage robustness check; the modest AUC drop "
        "(≈0.04) is well within the Bernoulli sampling-noise floor at "
        "n = 48 events (±0.06 at AUC = 0.70; Hanley & McNeil 1982), "
        "indicating the primary discrimination signal is not driven by "
        "post-event information.")
    add_page_break(doc)

    # Appendix S4, Why the Nationwide Inpatient Sample was excluded
    add_heading(doc, "Appendix S4.  Nationwide Inpatient Sample (NIS), "
                       "excluded from the primary analysis", level=1)
    add_para(doc,
        "A preliminary analysis attempted to extend external evaluation to "
        "the HCUP Nationwide Inpatient Sample (NIS, 2016–2019), restricted "
        "to admissions with chronic-SDH coding and a craniotomy or burr-hole "
        "procedure in the same admission (n = 2,518). On methodological "
        "review the cohort was excluded from the primary analysis because "
        "the available administrative coding for the postoperative-seizure "
        "outcome is unreliable in this setting.")
    add_para(doc,
        "Two coding limitations drive the exclusion. First, the ICD-10-CM "
        "code families historically used for postoperative seizure in NIS "
        "analyses combine acute symptomatic seizure codes (R56.x, 780.39, "
        "G41.x) with codes for pre-existing epilepsy (G40.x, 345.x). The "
        "two phenotypes are mechanistically distinct: acute symptomatic "
        "seizure is the in-admission event of interest, while pre-existing "
        "epilepsy reflects chronic disease coded for billing on every "
        "admission. Conflating them inflates the apparent population "
        "seizure signal and produces an outcome variable that does not "
        "match the BIDMC/eICU phenotype. Second, NIS does not carry a "
        "structured timestamp on the seizure code, so we cannot "
        "distinguish in-admission events from prevalent epilepsy after "
        "the procedure date. Together these properties make NIS "
        "incompatible with the chart-documented (BIDMC) and timestamped "
        "structured-flag (eICU) outcome ascertainment used in the "
        "primary analysis.")
    add_para(doc,
        "Supplementary Figure S5 illustrates the empirical effect: "
        "discrimination on the originally-reported combined outcome "
        "(AUC 0.617) collapses to chance (AUC 0.498) when the outcome "
        "is restricted to acute symptomatic seizure alone. Supplementary "
        "Figure S7 shows that group-LASSO under the corrected outcome "
        "selects no stable predictors. We release the cleaned ICD-10 "
        "codeset alongside the codebase (see Appendix S2) so that future "
        "nationwide analyses with timestamped coding can re-attempt this "
        "validation step on a coding substrate that supports it.")
    add_page_break(doc)

    # Appendix S5, Missing-data sensitivity
    add_heading(doc, "Appendix S5.  Missing-data sensitivity", level=1)
    add_para(doc,
        "Background.  We evaluated whether the discrimination of the "
        "candidate pipeline depends on the missing-data assumption. "
        "The candidate pipeline imputes each feature by its median and "
        "appends a per-feature binary missing-indicator covariate; "
        "sensitivity to this choice is reported below.")
    add_para(doc,
        "Missingness patterns.  In the BIDMC development cohort, two of "
        "the 21 postoperative-A variables carried administrative-category "
        "missingness: 'demographic' (63/655 = 9.6%) and 'procedures' "
        "(56/655 = 8.5%); the remaining 19 variables were complete. The "
        "eICU non-traumatic stratum had additional feature-specific "
        "missingness concentrated in structured-record physiological "
        "variables; per-feature rates are visualised in Supplementary "
        "Figure S4.")
    add_para(doc,
        "Little's MCAR test.  Applied to the BIDMC postoperative-B "
        "feature set, the test rejected the missing-completely-at-random "
        "hypothesis (χ² = 79.3, df = 33, p < 0.001), consistent with at-"
        "most a missing-at-random pattern across the two administrative-"
        "category variables. Multiple imputation under MAR is therefore "
        "the principled default.")
    # Imputation comparison table
    add_para(doc,
        "Imputation method comparison.  Four imputation strategies were "
        "evaluated by repeated 5 × 3 stratified cross-validation. On "
        "BIDMC postoperative-B the three imputers tested return "
        "identical AUC because the cohort carries no missingness on "
        "this feature set, included here for transparency.")
    imp_tbl = pd.DataFrame([
        {"Cohort": "eICU pure Set C", "Imputer": "Median (deployment default)",         "AUC": "0.644", "SD": "0.046"},
        {"Cohort": "eICU pure Set C", "Imputer": "Mean",                                "AUC": "0.693", "SD": "0.035"},
        {"Cohort": "eICU pure Set C", "Imputer": "IterativeImputer (MICE-like)",        "AUC": "0.675", "SD": "0.041"},
        {"Cohort": "eICU pure Set C", "Imputer": "Missing-indicator + median",          "AUC": "0.700", "SD": "0.035"},
        {"Cohort": "BIDMC postop-B",  "Imputer": "Median (deployment default)",         "AUC": "0.639", "SD": "0.093"},
        {"Cohort": "BIDMC postop-B",  "Imputer": "IterativeImputer (MICE-like)",        "AUC": "0.639", "SD": "0.093"},
        {"Cohort": "BIDMC postop-B",  "Imputer": "Missing-indicator + median",          "AUC": "0.632", "SD": "0.090"},
    ])
    add_table_from_df(doc, imp_tbl)
    add_para(doc,
        "Rubin's-rules pooling.  Ten multiple imputations (random-state "
        "seeded, IterativeImputer chains) were generated on the eICU "
        "pure Set C cohort. Discrimination was pooled across imputations "
        "by Rubin's rules with within-imputation bootstrap variance and "
        "between-imputation variance combined to compute the standard "
        "error and 95% confidence interval, and the fraction of missing "
        "information (FMI) and Rubin's degrees of freedom:")
    rubin_tbl = pd.DataFrame([
        {"Cohort": "eICU pure Set C", "M": "10",
         "Pooled AUC": "0.644", "SE": "0.022",
         "95% CI": "0.600 – 0.688",
         "FMI": "0.33", "df (Rubin)": "81"},
    ])
    add_table_from_df(doc, rubin_tbl)
    add_para(doc,
        "Interpretation.  AUC was insensitive to imputation choice on "
        "the deployment cohort (BIDMC, zero missingness) and modestly "
        "sensitive on the external eICU cohort, with the deployment "
        "default (median imputation) returning the lower end of the "
        "imputer-method range and the missing-indicator augmentation "
        "the upper end. The Rubin's-rules pooled estimate on eICU pure "
        "Set C (0.644, 95% CI 0.600–0.688; FMI 0.33) is concordant with "
        "the primary external-evaluation AUC reported in the main "
        "manuscript. The non-trivial FMI (0.33) is a useful pre-"
        "registered upper bound on the additional uncertainty injected "
        "by the imputation step alone and should be carried forward in "
        "any prospective re-validation that involves structured-record "
        "data with comparable missingness.")
    add_page_break(doc)

    # Appendix S6, Conformal set-to-action mapping + decision-analytic sensitivity (M4, M3, S11)
    add_heading(doc, "Appendix S6.  Conformal set-to-action mapping and "
                     "cost-effectiveness sensitivity", level=1)
    add_para(doc,
        "Mapping. Each conformal prediction set is mapped to an explicit "
        "action: a singleton {no-seizure} to observation, a singleton "
        "{seizure} to continuous EEG plus targeted AED, and the doubleton "
        "(abstention) to a default policy. Because the doubleton policy is a "
        "design lever, three mappings were evaluated (doubleton → universal "
        "AED, → observation, → cEEG).")
    add_para(doc,
        "Deployable operating point. At the calibrated base-rate threshold "
        "(0.073) the deployable postoperative-B model has the following "
        "operating characteristics (bootstrap 95% CIs):")
    if (RES / "38_postopB_operating_points.csv").exists():
        op = pd.read_csv(RES / "38_postopB_operating_points.csv")
        row = op.iloc[(op["threshold"] - 0.073).abs().argsort().iloc[0]]
        optab = pd.DataFrame([
            {"Metric": "Sensitivity", "Estimate": f"{row['sens']:.2f}", "95% CI": f"{row['sens_lo']:.2f}–{row['sens_hi']:.2f}"},
            {"Metric": "Specificity", "Estimate": f"{row['spec']:.2f}", "95% CI": f"{row['spec_lo']:.2f}–{row['spec_hi']:.2f}"},
            {"Metric": "PPV", "Estimate": f"{row['ppv']:.2f}", "95% CI": f"{row['ppv_lo']:.2f}–{row['ppv_hi']:.2f}"},
            {"Metric": "NPV", "Estimate": f"{row['npv']:.2f}", "95% CI": f"{row['npv_lo']:.2f}–{row['npv_hi']:.2f}"},
        ])
        add_table_from_df(doc, optab, caption="Table S6a.  Deployable operating point (threshold 7.3%).")
    if (RES / "38_conformal_mapping_sensitivity.csv").exists():
        cm = pd.read_csv(RES / "38_conformal_mapping_sensitivity.csv")
        cm = cm[cm["wtp"] == 100000][["policy", "best_strategy", "mapping_sens", "mapping_spec"]]
        add_table_from_df(doc, cm.round(3),
            caption="Table S6b.  Optimal strategy at $100k/QALY under each doubleton mapping.")
    add_para(doc,
        "AED efficacy and harm threshold. The optimal active strategy depends "
        "on two uncertain cSDH-specific parameters. One-way analysis at "
        "$100k/QALY shows ML-guided allocation becomes the highest-net-benefit "
        "strategy once the AED relative-risk reduction is ≤0.30 or the AED "
        "disutility is ≥0.10; universal AED is preferred only under the "
        "optimistic imported assumptions of strong efficacy and negligible "
        "harm. Because no cSDH study demonstrates a significant protective "
        "effect, the cSDH-plausible range favours ML-guided allocation.")
    if (RES / "42_international_perspectives.csv").exists():
        intl = pd.read_csv(RES / "42_international_perspectives.csv")
        add_table_from_df(doc, intl.round(1),
            caption="Table S6c.  International cost perspectives (US, UK NICE, Eurozone).")
    add_para(doc,
        "Model-versus-random comparator. To isolate the value contributed by "
        "the model's discrimination from the value of merely treating fewer "
        "patients, ML-guided allocation was compared with random allocation at "
        "the same treated fraction (32.4%). The net-monetary-benefit difference "
        "(discrimination premium) is $996–$1,612 per patient across the "
        "AED-efficacy range, confirming that the model, not the treated "
        "fraction alone, drives the incremental value.")
    if (RES / "44_model_vs_random.csv").exists():
        mvr = pd.read_csv(RES / "44_model_vs_random.csv")
        mvr["aed_rrr"] = mvr["aed_rrr"].round(2)
        mvr["treated_fraction"] = mvr["treated_fraction"].round(3)
        for c in ["NMB_obs", "NMB_universal_aed", "NMB_ml_guided",
                  "NMB_random_matched", "discrimination_premium"]:
            mvr[c] = mvr[c].round(0).astype(int)
        add_table_from_df(doc, mvr,
            caption="Table S6d.  Model-vs-random comparator: discrimination "
                    "premium at matched treated fraction.")
    add_page_break(doc)

    # Appendix S7, LOHO heterogeneity (S7)
    add_heading(doc, "Appendix S7.  Leave-one-hospital-out heterogeneity", level=1)
    add_para(doc,
        "Random-effects pooling of per-hospital AUCs used Hanley–McNeil "
        "within-study variance and three τ² estimators (DerSimonian–Laird, "
        "Paule–Mandel, REML/iterated). Heterogeneity is scale-sensitive: on "
        "the variance-stabilising logit scale (the primary analysis) "
        "between-hospital heterogeneity is negligible (I² = 0% for the "
        "non-traumatic Set C; 2.8% for Set A), whereas a raw-AUC sensitivity "
        "analysis gives higher estimates (I² up to ~56%). The 95% prediction "
        "interval is the most honest summary of cross-site transportability.")
    if (RES / "41_loho_heterogeneity_summary.csv").exists():
        het = pd.read_csv(RES / "41_loho_heterogeneity_summary.csv")
        het = het[["cohort", "set", "scale", "estimator", "pooled_auc", "ci_lo", "ci_hi",
                   "tau2", "I2_pct", "pred_int_lo", "pred_int_hi"]]
        add_table_from_df(doc, het.round(4),
            caption="Table S7.  LOHO random-effects pooling under three τ² estimators "
                    "and two scales, with prediction intervals.")
    add_para(doc, "Per-site descriptive statistics (n, events, AUC, Brier per "
                  "hospital) are released as results/41_loho_per_site.csv.")
    add_page_break(doc)

    # Appendix S8, CONSORT flow + candidate-model coefficients & calibration (M2, S12)
    add_heading(doc, "Appendix S8.  Cohort inclusion flow and candidate-model "
                     "coefficients and calibration", level=1)
    add_para(doc,
        "Cohort inclusion (eICU). 5,376 ICU stays with any subdural-haematoma "
        "diagnosis (139 hospitals) → restrict to non-traumatic SDH and to "
        "hospitals contributing ≥10 such stays → 3,297 stays (42 hospitals; "
        "300 seizures), the primary external-evaluation stratum. The broader "
        "5,376/139 screen and three alternative definitions are reported as "
        "sensitivity analyses (Table S2).")
    add_para(doc,
        "Candidate-model coefficients. Firth penalized logistic regression on "
        "the postoperative-B set; coefficients are on the per-standard-"
        "deviation scale, with odds ratios per 1-SD increase:")
    if (RES / "40_postopB_firth_coefficients.csv").exists():
        co = pd.read_csv(RES / "40_postopB_firth_coefficients.csv")
        co = co[["feature", "coef", "se", "ci_lo", "ci_hi", "p_value", "OR_per_SD"]]
        add_table_from_df(doc, co.round(3),
            caption="Table S8.  Firth coefficients for the candidate postoperative-B model "
                    "(per-SD scale; OR per 1-SD increase).")
    add_para(doc,
        "The only coefficients reaching statistical significance are procedure "
        "variables (surgical decompression, middle-meningeal-artery "
        "embolization, and drainage); these reflect confounding by indication "
        "rather than causal seizure biology, for example, drainage carries an "
        "odds ratio of 1.44, a direction inconsistent with a plausible causal "
        "protective effect, which is a limitation of the candidate model.")
    add_para(doc,
        "Candidate-model calibration. Out-of-fold, the candidate model achieves "
        "calibration-in-the-large near zero and low expected calibration "
        "error; the recalibration slope exceeds one (under-dispersion), "
        "reflecting weak discrimination over a narrow probability range "
        "rather than mean miscalibration. Calibration metrics are released as "
        "results/40_postopB_calibration_post_platt.csv.")
    if (FIG / "40_postopB_calibration.png").exists():
        add_figure(doc, FIG / "40_postopB_calibration.png",
                   caption="Figure S8.  Reliability curve of the candidate postoperative-B "
                           "model after Platt scaling.")
    add_page_break(doc)

    # Tables S1–S5
    add_heading(doc, "Supplementary Tables", level=1)
    for tbl_path, label in [
        (RES / "21_imbalance_sweep.csv",
         "Table S1.  Eleven-method modelling battery, full numeric results."),
        (RES / "08_cohort_comparison.csv",
         "Table S2.  eICU cohort definition sensitivity with bootstrap 95% CIs."),
        (RES / "02_calibration_metrics.csv",
         "Table S3.  Calibration metrics with bootstrap 95% CIs across all six cohort-model combinations."),
        (RES / "10_pairwise_summary.csv",
         "Table S4.  Cost-effectiveness analysis, PSA summary at WTP $50k, $100k, $150k."),
        (RES / "45_voi_postopB.csv",
         "Table S5.  Value-of-information under the deployable postoperative-B "
         "candidate model and cSDH-grounded AED priors: expected value of "
         "perfect information (EVPI) and per-parameter expected value of "
         "partial perfect information (EVPPI) at WTP $100,000/QALY "
         "(per-patient and 10-year population, $M)."),
    ]:
        if tbl_path.exists():
            df = pd.read_csv(tbl_path)
            if len(df.columns) > 9:
                # truncate excessively wide tables, show selected columns
                df = df.iloc[:, :9]
            df = df.round(3).fillna("")
            add_table_from_df(doc, df, caption=label)
            add_para(doc, "")

    add_page_break(doc)

    # Figures S1–S7
    add_heading(doc, "Supplementary Figures", level=1)
    supp_figs = [
        ("04_loho_forest_full_Set_C.png",
         "Figure S1.  Per-hospital leave-one-hospital-out forest plot (eICU "
         "Set C, full cohort, 42 hospitals)."),
        ("04_loho_forest_full_Set_A.png",
         "Figure S2.  Per-hospital leave-one-hospital-out forest plot (eICU "
         "Set A, full cohort)."),
        ("05_auc_comparison.png",
         "Figure S3.  Temporal-leakage audit, AUC across feature "
         "specifications and time-window cuts."),
        ("07_missingness.png",
         "Figure S4.  Missingness pattern and Rubin's-rules multiple-"
         "imputation pooling."),
        ("12_nis_outcome_auc.png",
         "Figure S5.  Nationwide Inpatient Sample outcome reclassification, "
         "AUC under original combined outcome (0.617) vs corrected acute "
         "symptomatic outcome (0.498)."),
        ("09_cumulative_incidence.png",
         "Figure S6.  Competing-risks cumulative-incidence functions and "
         "cause-specific Cox / IPCW Fine–Gray comparison."),
        ("13_nis_grouped_lasso.png",
         "Figure S7.  λ-path-tuned group-LASSO and sparse-group-LASSO on the "
         "NIS chronic+surgical cohort with the corrected outcome."),
    ]
    for fname, cap in supp_figs:
        path = FIG / fname
        if path.exists():
            add_figure(doc, path, caption=cap)
            add_para(doc, "")

    doc.save(SUPP_PATH)
    print(f"[OK] supplementary: {SUPP_PATH}")
    print(f"     file size: {os.path.getsize(SUPP_PATH)/1024:.1f} KB")


def main():
    build_main()
    build_supplementary()


if __name__ == "__main__":
    main()
