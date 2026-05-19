"""Task 32 — Build the Code Companion document.

A long-form, educational walkthrough of the entire codebase, written for a
reader who is comfortable reading clinical research papers but who is new to
Python and to machine-learning code. Every key script is annotated with a
short narrative ('what this script does, in plain English'), the code listing
itself, and a 'line-by-line' guide that explains each construct.

The companion also includes:
  • a Python primer covering the constructs that recur in our scripts
  • a deep dive on the static HTML dashboards
  • an explanation of the module callgraph
  • a list of Claude skills the reader can use to study the code further

Output: github_repo/docs/Code_Companion.docx
"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _shared import OUT

REPO = OUT / "github_repo"
DOCS = REPO / "docs"
SCRIPTS = REPO / "scripts"
SITE = REPO / "site"
OUT_PATH = DOCS / "Code_Companion.docx"

# ─── style helpers ──────────────────────────────────────────
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x14, 0x1F, 0x3A)
        r.font.name = "Calibri"
    return h

def add_para(doc, text, *, italic=False, bold=False, size=11,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False,
             space_after=8):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(space_after)
    if indent: p.paragraph_format.first_line_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Calibri"
    r.italic = italic; r.bold = bold
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.name = "Calibri"
    return p

def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_code_block(doc, code, language="python", max_lines=None):
    """Render a code block in a single-cell table with monospace font,
    light-grey background, and line numbers."""
    lines = code.splitlines()
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines] + [
            f"  ... ({len(code.splitlines()) - max_lines} more lines elided "
            f"— see full script in scripts/) ..."
        ]
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    _shade_cell(cell, "F4F4F4")
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.15
    for i, line in enumerate(lines, 1):
        r = p.add_run(f"{i:>3} │ {line}\n")
        r.font.name = "Menlo"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    # Remove trailing blank line
    p.runs[-1].text = p.runs[-1].text.rstrip("\n")
    add_para(doc, "", space_after=2)

def add_callout(doc, label, body, color="#FFF7E0", border="#AA7700"):
    """Highlighted box for 'why this matters' notes."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade_cell(cell, color.lstrip("#"))
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"
    r2 = p.add_run(body)
    r2.font.size = Pt(11); r2.font.name = "Calibri"
    add_para(doc, "", space_after=2)

def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light List Accent 1"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; c.text = str(h)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10); r.font.name = "Calibri"; r.bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10); r.font.name = "Calibri"
    add_para(doc, "", space_after=4)

def page_break(doc): doc.add_page_break()

def setup_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0); section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)

def read_script(name):
    p = SCRIPTS / name
    if not p.exists(): return f"# {name} not found"
    return p.read_text(encoding="utf-8")

def read_site(name):
    p = SITE / name
    if not p.exists(): return f"# {name} not found"
    return p.read_text(encoding="utf-8")

# ─── document body ─────────────────────────────────────────
def build():
    doc = Document()
    setup_document(doc)

    # ── Title ──
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Code Companion")
    tr.bold = True; tr.font.size = Pt(28); tr.font.name = "Calibri"
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("A line-by-line guide to the csdh-postop-seizure-risk analysis code")
    tr.italic = True; tr.font.size = Pt(15); tr.font.name = "Calibri"
    add_para(doc, "")
    add_para(doc, "Niels Pacheco-Barrios", alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Department of Neurosurgery, Beth Israel Deaconess Medical Center, "
                   "Harvard Medical School", alignment=WD_ALIGN_PARAGRAPH.CENTER,
              size=10, italic=True)
    add_para(doc, "")
    add_callout(doc, "How to read this document.",
                 "Each chapter explains one part of the codebase in plain "
                 "English, then shows the actual code with line numbers, and "
                 "then unpacks what each block is doing. You do not need to "
                 "be fluent in Python to follow along; the Python primer in "
                 "Chapter 1 introduces the constructs that recur throughout. "
                 "If you only have an hour, read Chapter 1, Chapter 2 "
                 "(_shared.py), and Chapter 6 (Firth + conformal). Those "
                 "three pieces are the spine of the analysis.")
    page_break(doc)

    # ── Table of contents ──
    add_heading(doc, "Contents", level=1)
    toc = [
        ("Chapter 1", "A short Python primer for clinicians", ""),
        ("Chapter 2", "_shared.py — the foundations every script depends on", ""),
        ("Chapter 3", "02_calibration.py — measuring how trustworthy the probabilities are", ""),
        ("Chapter 4", "04_loho.py — leave-one-hospital-out validation", ""),
        ("Chapter 5", "14_decision_tree.py — the TreeAge-style cost-effectiveness model", ""),
        ("Chapter 6", "24_firth_bayes_lr.py — the deployment model", ""),
        ("Chapter 7", "25_conformal_prediction.py — turning probabilities into bedside decisions", ""),
        ("Chapter 8", "16_voi_evpi.py — value-of-information analysis", ""),
        ("Chapter 9", "29_main_figures_jnnp.py — publication-grade figures", ""),
        ("Chapter 10", "27_build_jnnp_manuscript.py — building the Word manuscript from code", ""),
        ("Chapter 11", "30_export_calculator_assets.py — packaging the model for the dashboard", ""),
        ("Chapter 12", "The interactive dashboards — HTML, CSS, JavaScript", ""),
        ("Chapter 13", "The callgraph — how the scripts depend on each other", ""),
        ("Chapter 14", "Other scripts — what they do, in one paragraph each", ""),
        ("Chapter 15", "Further study — Claude skills that help you learn the code", ""),
    ]
    add_table(doc, ["Chapter", "Title", ""], toc)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 1 — Python primer
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 1.  A short Python primer for clinicians", level=1)
    add_para(doc,
        "Python is a programming language often described as 'executable "
        "pseudocode' because the constructs in a Python script tend to "
        "translate cleanly into English sentences. This chapter introduces "
        "the small set of building blocks that recur throughout our analysis "
        "code. If a line of code in any later chapter ever confuses you, "
        "the explanation is almost certainly in this chapter.", indent=True)

    add_heading(doc, "1.1  How a Python file is organised", level=2)
    add_para(doc,
        "A Python file (extension .py) is read from top to bottom. The "
        "topmost few lines usually do two things in order: state the "
        "purpose of the file in a triple-quoted block (a 'docstring'), "
        "then import the libraries the file will need.", indent=True)
    add_code_block(doc,
        '"""Task 24 — Firth penalized + Bayesian logistic regression.\n'
        'A docstring describes what this file does, who reads it, and what\n'
        'it produces. Triple double-quotes mark a multi-line string.\n'
        '"""\n'
        'import os, sys\n'
        'import numpy as np, pandas as pd\n'
        'from sklearn.linear_model import LogisticRegression\n'
        'from _shared import load_bidmc, RES, FIG, SEED\n')
    add_para(doc,
        "Lines starting with the word import bring in code from elsewhere. "
        "import numpy as np gives us a shorthand: from this line on, np "
        "stands for the numpy library. from sklearn.linear_model import "
        "LogisticRegression pulls one specific tool out of a larger "
        "library — exactly like importing one drug class from a "
        "pharmacopoeia.", indent=True)

    add_heading(doc, "1.2  Variables, lists, and dictionaries", level=2)
    add_code_block(doc,
        '# A single value\n'
        'age = 73\n'
        'prevalence = 0.073\n'
        'cohort_name = "BIDMC"\n'
        '\n'
        '# A list — an ordered collection\n'
        'features = ["age", "sex", "preop_gcs", "sdh_thickness"]\n'
        'features.append("mid_shift")\n'
        '\n'
        '# A dictionary — a mapping from keys to values\n'
        'feature_means = {"age": 73, "sex": 0.68, "preop_gcs": 14.1}\n'
        'feature_means["epilepsy_hx"] = 0.02\n')
    add_para(doc,
        "Variables are labels that point to values. Lists hold ordered "
        "collections and use square brackets. Dictionaries map a key to a "
        "value and use curly braces. The hash sign # introduces a comment "
        "that the interpreter ignores — comments are notes for human "
        "readers.", indent=True)

    add_heading(doc, "1.3  Functions: code with a name and inputs", level=2)
    add_code_block(doc,
        'def bootstrap_auc(y, p, n_boot=1000, seed=42):\n'
        '    """Compute the AUC and a 95% bootstrap confidence interval."""\n'
        '    rng = np.random.default_rng(seed)\n'
        '    bs = []\n'
        '    for _ in range(n_boot):\n'
        '        idx = rng.integers(0, len(y), len(y))\n'
        '        bs.append(roc_auc_score(y[idx], p[idx]))\n'
        '    lo, hi = np.percentile(bs, [2.5, 97.5])\n'
        '    return float(roc_auc_score(y, p)), float(lo), float(hi)\n')
    add_para(doc,
        "A function is a named recipe. def starts the definition, the name "
        "follows (here bootstrap_auc), and the inputs go inside the "
        "parentheses. Inputs can have default values (n_boot=1000 means "
        "'use 1000 unless the caller specifies otherwise'). The indented "
        "block under def is the body. The for line is a loop: it runs the "
        "indented code 1000 times. The variable bs starts as an empty list "
        "and grows by one entry each iteration. The return line at the end "
        "is what the function gives back to whoever called it.", indent=True)

    add_heading(doc, "1.4  Pandas, numpy, and scikit-learn — the three "
                       "libraries that do the heavy lifting", level=2)
    add_para(doc,
        "Three external libraries do almost all of the analytical work in "
        "this codebase. numpy provides fast numerical arrays — think of it "
        "as Excel without the user interface. pandas provides DataFrames, "
        "which are tables with named columns and labelled rows; pandas is "
        "what we use to load CSVs, filter rows, and compute summaries. "
        "scikit-learn (imported as sklearn) provides the machine-learning "
        "models, cross-validation splitters, and the standardisation "
        "machinery (Pipeline, ColumnTransformer, StandardScaler).", indent=True)

    add_heading(doc, "1.5  Pipelines and the scikit-learn fit / predict pattern", level=2)
    add_code_block(doc,
        'pipe = Pipeline([\n'
        '    ("imputer", SimpleImputer(strategy="median")),\n'
        '    ("scaler",  StandardScaler()),\n'
        '    ("classifier", FirthLogisticRegression()),\n'
        '])\n'
        'pipe.fit(X_train, y_train)\n'
        'probs = pipe.predict_proba(X_test)[:, 1]\n')
    add_para(doc,
        "A scikit-learn pipeline is a chain of steps. The list of tuples "
        "names each step and tells the pipeline what to do at that step. "
        "When we call pipe.fit, the pipeline runs imputation on the "
        "training data, then standardisation, then trains the classifier. "
        "When we call pipe.predict_proba, the pipeline runs the same "
        "imputation and standardisation (using the training-time parameters) "
        "and then asks the classifier for predicted probabilities. The "
        "[:, 1] at the end is shorthand for 'every row, column index 1' — "
        "i.e., the probability of the positive (seizure) class.", indent=True)
    add_callout(doc, "Why pipelines matter.",
                "A pipeline guarantees that the same preprocessing applied "
                "to the training data is applied — using the training-time "
                "statistics — to any new data. Without a pipeline it is "
                "easy to compute the column means on the test set and "
                "introduce a leakage bias.")

    add_heading(doc, "1.6  Cross-validation in three lines", level=2)
    add_code_block(doc,
        'rskf = RepeatedStratifiedKFold(n_splits=5, n_repeats=5, random_state=42)\n'
        'for train_idx, test_idx in rskf.split(X, y):\n'
        '    pipe.fit(X.iloc[train_idx], y.iloc[train_idx])\n'
        '    p_oof[test_idx] = pipe.predict_proba(X.iloc[test_idx])[:, 1]\n')
    add_para(doc,
        "RepeatedStratifiedKFold is a class that, when iterated over with "
        "for, yields pairs of (train_idx, test_idx). Each pair is one fold "
        "of cross-validation: train on 80% of the data, evaluate on the "
        "held-out 20%, save the predicted probabilities for the held-out "
        "patients into a vector called p_oof (out-of-fold). After all "
        "folds are processed, p_oof contains one prediction for every "
        "patient — each generated by a model that did not see that "
        "patient at training time.", indent=True)

    add_heading(doc, "1.7  Where files live", level=2)
    add_para(doc,
        "All our scripts assume a fixed folder layout. Inputs (the CSVs of "
        "patient data) sit at the project root. The scripts themselves "
        "live in scripts/. Outputs go into results/ (CSV tables), figures/ "
        "(PNG and PDF plots) and cache/ (intermediate prediction arrays "
        "the next script might want to reuse). The helper module "
        "_shared.py defines these locations once so every other script "
        "imports them — Chapter 2 covers that file in detail.", indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 2 — _shared.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 2.  _shared.py — the foundations every script depends on", level=1)
    add_para(doc,
        "_shared.py is the file every other script in the project imports. "
        "It defines the path constants that say where data lives and where "
        "outputs go, the feature lists that name the variables used in "
        "each model specification, the loader functions that read the raw "
        "CSVs and apply the standard filtering rules, and the pipeline "
        "factories that build a fresh scikit-learn pipeline on demand.",
        indent=True)
    add_para(doc,
        "Reading _shared.py is the single highest-yield investment of your "
        "time, because the patterns it establishes appear in every other "
        "script. We will walk through it section by section.", indent=True)

    src_shared = read_script("_shared.py")
    add_heading(doc, "2.1  Imports and global thread caps", level=2)
    add_para(doc,
        "The file opens with environment-variable settings that force the "
        "underlying numerical libraries (OpenBLAS, MKL, OpenMP) to use a "
        "single CPU thread per process. On Apple Silicon, multi-threaded "
        "scikit-learn can deadlock; setting n_jobs=1 throughout the "
        "codebase prevents that.", indent=True)
    add_code_block(doc, "\n".join(src_shared.splitlines()[:23]))

    add_heading(doc, "2.2  Path constants", level=2)
    add_code_block(doc, "\n".join(src_shared.splitlines()[12:23]))
    add_para(doc,
        "ROOT is the project's top-level folder. OUT is revision_analyses/, "
        "RES is its results/ subfolder, FIG is figures/, CACHE is for "
        "intermediate arrays. The for loop at the bottom creates each of "
        "those folders if they do not yet exist (parents=True means "
        "'create any missing parent folders too'; exist_ok=True means "
        "'do not raise an error if it already exists').", indent=True)

    add_heading(doc, "2.3  Feature lists", level=2)
    add_para(doc,
        "POSTOP_A_FEATURES is a Python list of 21 column names taken from "
        "the BIDMC CSV. Defining it here, once, means every script that "
        "fits the postoperative-A model picks up exactly the same 21 "
        "predictors. Changing the list in one place updates all downstream "
        "scripts. POSTOP_B_FEATURES is the same list with the three "
        "'leakage-suspect' variables removed — that is the temporal-"
        "leakage robustness analysis.", indent=True)
    # Show the postop_A feature list section
    add_code_block(doc, "\n".join(src_shared.splitlines()[24:46]))

    add_heading(doc, "2.4  The loader functions", level=2)
    add_code_block(doc, "\n".join(src_shared.splitlines()[47:62]))
    add_para(doc,
        "load_bidmc reads the cleaned BIDMC CSV and applies a single "
        "recoding: aed_timing values of 2 (meaning 'unknown timing') are "
        "collapsed into 1, so the variable becomes a clean binary "
        "indicator. load_eicu does the same for eICU; load_eicu_pure "
        "additionally filters to patients with no prior seizure history, "
        "no pre-admission AED, and no mechanical ventilation — the 'pure "
        "post-craniotomy' sensitivity cohort.", indent=True)

    add_heading(doc, "2.5  Pipeline factories", level=2)
    add_para(doc,
        "Each pipeline factory is a small function that returns a fresh, "
        "untrained scikit-learn Pipeline. Returning a fresh pipeline (rather "
        "than reusing one) is important because pipelines are mutable: "
        "fitting one in a loop would carry over state from the previous "
        "iteration. The factories let every CV fold start from a clean "
        "slate. Below is the postop_A factory.", indent=True)
    # Find make_pipeline_postopA in the source
    idx = src_shared.find("def make_pipeline_postopA")
    end = src_shared.find("\ndef ", idx + 1)
    add_code_block(doc, src_shared[idx:end].rstrip())
    add_para(doc,
        "ColumnTransformer holds a list of (name, transformer, columns) "
        "triples. Here the only transformer is the imputation+scaling "
        "sub-pipeline, applied to every feature in POSTOP_A_FEATURES. "
        "BalancedRandomForestClassifier is the imblearn variant of "
        "scikit-learn's RandomForestClassifier; it under-samples the "
        "majority class inside each bootstrap iteration so the trees see "
        "balanced data. n_estimators=300 builds 300 trees; n_jobs=N_JOBS "
        "is 1 (set above); random_state=SEED is 42 for reproducibility.",
        indent=True)

    add_heading(doc, "2.6  Out-of-fold prediction helper", level=2)
    idx = src_shared.find("def oof_predictions")
    end = src_shared.find("\ndef ", idx + 1)
    add_code_block(doc, src_shared[idx:end].rstrip())
    add_para(doc,
        "This function returns one predicted probability per patient by "
        "running repeated stratified K-fold cross-validation and averaging "
        "the predictions whenever a patient appears in multiple test "
        "folds (which happens because n_repeats > 1). The accumulator "
        "trick — adding into p_acc and incrementing n_acc — is the "
        "standard idiom for averaging without storing every individual "
        "prediction.", indent=True)
    add_callout(doc, "Why averaged OOF predictions.",
                "Averaged out-of-fold predictions are what calibration "
                "metrics, decision-curve analyses, and conformal layers "
                "consume downstream. Computing them once in this helper "
                "and caching them to disk means every downstream script "
                "operates on the same reproducible prediction vector.")
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 3 — 02_calibration.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 3.  02_calibration.py — measuring how trustworthy "
                       "the probabilities are", level=1)
    add_para(doc,
        "A model's AUC measures rank-order discrimination — does the model "
        "assign higher probabilities to patients who go on to seize than "
        "to those who do not? A model's calibration measures whether those "
        "probabilities are themselves trustworthy — when the model says "
        "'10%', do roughly 10% of those patients actually seize? Calibration "
        "is the metric that shapes clinical decisions. This script computes "
        "it.", indent=True)
    src = read_script("02_calibration.py")
    add_code_block(doc, src, max_lines=120)
    add_para(doc,
        "The script defines run_one — a small wrapper that takes a model "
        "factory, the X / y data, and the CV specification, and returns "
        "the calibration metrics plus the OOF predictions. The metrics "
        "themselves are computed by calibration_metrics in _shared.py: "
        "Brier score (the mean squared error between predicted "
        "probability and actual outcome), calibration-in-the-large (the "
        "average predicted probability minus the actual event rate), "
        "calibration slope (how much the predicted probabilities need to "
        "be rescaled to match observed), and calibration intercept (any "
        "constant shift). A well-calibrated model has Brier near "
        "p×(1−p), CITL near zero, slope near one, and intercept near "
        "zero.", indent=True)
    add_callout(doc, "Reading the calibration figure.",
                "Each model gets a panel in Figure 2 showing the "
                "calibration curve (observed event rate vs predicted "
                "probability in deciles). The diagonal is perfect "
                "calibration. The histogram at the bottom of each panel "
                "shows how predicted probabilities are distributed; a "
                "model that only ever predicts close to the base rate is "
                "uninformative no matter how its calibration line looks.")
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 4 — 04_loho.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 4.  04_loho.py — leave-one-hospital-out validation", level=1)
    add_para(doc,
        "External validation in eICU is more demanding than ordinary "
        "cross-validation because it asks: does the model generalise "
        "from hospital A to hospital B that the model has never seen? "
        "We answer that by leaving one hospital out at a time, training "
        "on the remaining 138 hospitals, and scoring the held-out "
        "hospital. With 139 hospitals this means 139 separate model fits.",
        indent=True)
    src = read_script("04_loho.py")
    add_code_block(doc, src, max_lines=150)
    add_para(doc,
        "Two helpers do the real work. loho_for handles the leave-one-"
        "hospital-out loop and writes a CSV row per hospital. "
        "random_effects_pool turns those per-hospital AUCs into a single "
        "meta-analytic estimate using the DerSimonian–Laird method "
        "(originally developed to combine treatment effects across "
        "randomised trials, but mathematically applicable to any set of "
        "study-level estimates with within-study variance). Hanley and "
        "McNeil's 1982 formula gives us the within-study variance of an "
        "AUC; together they produce the random-effects pooled AUC and "
        "the τ² between-hospital heterogeneity reported in Figure 1 panel "
        "B.", indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 5 — 14_decision_tree.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 5.  14_decision_tree.py — the TreeAge-style "
                       "cost-effectiveness model", level=1)
    add_para(doc,
        "Decision-tree models in health economics are mathematical "
        "objects that compute the expected cost and expected QALYs of a "
        "strategy by enumerating every possible patient trajectory, "
        "weighting each trajectory by its probability, and adding the "
        "results. This script implements such a tree for the four "
        "strategies (observation, universal AED, ML-guided AED, "
        "ML-guided cEEG) and then renders it as Figure 4 in the "
        "manuscript.", indent=True)
    src = read_script("14_decision_tree.py")
    add_code_block(doc, src, max_lines=160)
    add_para(doc,
        "rollback_strategy returns the list of terminal leaves for a "
        "given strategy. Each leaf carries (path label, probability, "
        "cost, QALY). compute_ev does the rollback: it multiplies each "
        "leaf's probability by its cost (or QALY) and sums. render_tree "
        "draws the tree in matplotlib using the TreeAge visual "
        "convention — blue square for the decision node, green circle "
        "for chance nodes, red triangle for terminals — and writes the "
        "rolled-up E[Cost] and E[QALY] beside each strategy.", indent=True)
    add_callout(doc, "Where the parameter values come from.",
                "The probabilities, costs and utilities used in the "
                "rollback are not hard-coded into 14_decision_tree.py. "
                "They are imported (via the P dictionary at the top of the "
                "file) and ultimately trace back to the probabilistic "
                "sensitivity analysis in 10_11_cea_pairwise.py, which "
                "draws each parameter from its own beta or gamma "
                "distribution. The point estimates used here are the "
                "means of those distributions.")
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 6 — 24_firth_bayes_lr.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 6.  24_firth_bayes_lr.py — the deployment model", level=1)
    add_para(doc,
        "This is the script that produces the manuscript's deployment "
        "model: Firth penalized logistic regression. With 48 events in "
        "the BIDMC cohort, ordinary maximum-likelihood logistic regression "
        "is unstable; Firth's penalty (a small bias correction derived "
        "from the Jeffreys prior on the likelihood) stabilises the "
        "coefficient estimates and produces valid confidence intervals "
        "even with rare events.", indent=True)
    src = read_script("24_firth_bayes_lr.py")
    add_code_block(doc, src, max_lines=180)
    add_para(doc,
        "The class BayesianLogReg is a from-scratch implementation of "
        "logistic regression with Gaussian priors centred at user-"
        "supplied means rather than at zero. This is the machinery for "
        "transfer learning from the eICU coefficient estimates. The "
        "Newton-Raphson loop in BayesianLogReg.fit alternates between "
        "computing the gradient and Hessian of the penalised negative "
        "log-likelihood and taking a Newton step until convergence. "
        "derive_eicu_priors fits an elastic-net LR on eICU's shared "
        "features and returns those coefficients as the prior means.",
        indent=True)
    add_para(doc,
        "The headline finding produced by this script is that Firth "
        "matches the published BalancedRandomForest baseline on AUC "
        "(0.681 vs 0.676; DeLong p = 0.81) with a 3.3-fold improvement "
        "in Brier score. The Bayesian-with-eICU-priors variant, by "
        "contrast, degrades discrimination significantly — a finding "
        "that the manuscript interprets as a biological-mismatch "
        "warning, not a statistical failure.", indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 7 — 25_conformal_prediction.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 7.  25_conformal_prediction.py — turning "
                       "probabilities into bedside decisions", level=1)
    add_para(doc,
        "Conformal prediction is a recent framework that turns any "
        "classifier into a procedure that emits prediction sets with "
        "distribution-free coverage guarantees. The class-conditional "
        "variant we use (Mondrian conformal) guarantees coverage "
        "separately for each true class: when the true outcome is "
        "'seizure', the prediction set will include 'seizure' with at "
        "least 1−α probability; same for 'no seizure'.", indent=True)
    src = read_script("25_conformal_prediction.py")
    add_code_block(doc, src, max_lines=160)
    add_para(doc,
        "class_conditional_conformal does the work in eight lines: "
        "compute the nonconformity score 1 − P(true class) on every "
        "calibration example, take the (1−α)(n+1)/n empirical quantile "
        "of those scores separately for the positive and negative "
        "classes, and for each test point include each class in the "
        "prediction set if and only if its nonconformity does not exceed "
        "that class's quantile. The result is a set of one or two "
        "elements: a 'rule-out' singleton if only 'no seizure' is "
        "included, a 'rule-in' singleton if only 'seizure' is included, "
        "or a doubleton if both are included.", indent=True)
    add_callout(doc, "Why rule-out matters clinically.",
                "Most patients in any cSDH cohort do not seize. A model "
                "that can rule out seizure with a 90% coverage guarantee "
                "in a quarter of patients lets the clinician confidently "
                "omit AED prophylaxis in that quarter — exactly the "
                "subgroup most likely to be harmed by AED side effects.")
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 8 — 16_voi_evpi.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 8.  16_voi_evpi.py — value-of-information analysis", level=1)
    add_para(doc,
        "Value-of-information analysis answers a sharp question: of all "
        "the uncertain parameters in our decision model, which would "
        "actually change the optimal strategy if we knew them perfectly? "
        "EVPI is the population-scale upper bound on the value of any "
        "future study. EVPPI breaks that upper bound down per parameter, "
        "ranking research priorities.", indent=True)
    src = read_script("16_voi_evpi.py")
    add_code_block(doc, src, max_lines=170)
    add_para(doc,
        "compute_evpi is the population-EVPI formula: take the per-"
        "iteration maximum NMB across strategies (what we could earn if "
        "we always picked the right strategy in each Monte Carlo "
        "iteration), subtract the maximum mean NMB (what we earn by "
        "always picking the strategy that is best on average). The "
        "difference is the expected value of perfect information.", indent=True)
    add_para(doc,
        "evppi_strong_oakley implements the Strong–Oakley (2014) "
        "non-parametric regression approximation for per-parameter "
        "EVPPI. For each focal parameter φ, it fits a smoothing spline "
        "of NMB on φ per strategy, takes the maximum across strategies "
        "at each iteration, averages, and subtracts the same baseline "
        "as before. The output is a tornado chart (Figure 6 panel A) "
        "showing which parameters dominate the research-value frontier.",
        indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 9 — 29_main_figures_jnnp.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 9.  29_main_figures_jnnp.py — publication-grade figures", level=1)
    add_para(doc,
        "Every main paper figure is built by a dedicated function in "
        "this script (figure_1 … figure_6). The script enforces a "
        "consistent house style: Helvetica sans-serif, open-box axes "
        "(top and right spines removed), bold uppercase panel labels in "
        "the top-left of each panel, a restrained palette suitable for "
        "grayscale printing, 300-dpi PNG plus editable vector PDF "
        "output.", indent=True)
    src = read_script("29_main_figures_jnnp.py")
    add_code_block(doc, src, max_lines=120)
    add_para(doc,
        "The plt.rcParams.update block at the top of the file sets the "
        "house style once; every panel in every figure inherits from it "
        "automatically. The helper add_panel_label puts the bold A / B / "
        "C labels in the top-left of each panel; style_axis applies the "
        "tick direction and the y-only gridline; figure_legend_below "
        "places a single shared legend beneath the figure so legends "
        "never overlap data. Each figure function reads result CSVs, "
        "draws its panels, and saves both PNG and PDF.", indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 10 — 27_build_jnnp_manuscript.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 10.  27_build_jnnp_manuscript.py — building the "
                       "Word manuscript from code", level=1)
    add_para(doc,
        "This script writes the manuscript .docx file end-to-end using "
        "python-docx. Building the manuscript programmatically (rather "
        "than typing it in Word) means every numerical claim in the text "
        "is traceable to the results CSV that produced it, and every "
        "version of the manuscript is a deterministic function of the "
        "underlying analysis.", indent=True)
    add_para(doc,
        "The structure mirrors a normal Word workflow: setup_document "
        "sets page margins and the Normal style; add_heading adds "
        "headings; add_para adds paragraphs; add_runs handles paragraphs "
        "that mix regular text with superscript references; register_"
        "figure and register_table accumulate figures and tables into a "
        "registry; render_collected_tables_and_figures emits them all "
        "in a single 'Tables and Figures' section at the end. "
        "build_main glues everything together and saves the document.",
        indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 11 — 30_export_calculator_assets.py
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 11.  30_export_calculator_assets.py — packaging "
                       "the model for the dashboard", level=1)
    add_para(doc,
        "The browser-side calculator at site/calculator.html cannot run "
        "scikit-learn — it is JavaScript. To let the browser compute "
        "the same prediction the Python model would, this script extracts "
        "the deployment-ready coefficients, the feature scaler statistics, "
        "and the class-conditional conformal quantiles and writes them to "
        "site/model_assets.json. The browser then loads that JSON and "
        "applies a dot product and sigmoid in pure JavaScript to recover "
        "the model's prediction.", indent=True)
    src = read_script("30_export_calculator_assets.py")
    add_code_block(doc, src, max_lines=120)
    add_para(doc,
        "The Firth model is fit on the full BIDMC cohort (no held-out "
        "split — this is the deployment fit, not the evaluation fit). "
        "The Newton-Raphson loop inside FirthLogisticRegression "
        "converges in a few iterations. The class-conditional "
        "conformal quantiles are computed by repeated 5-fold splits with "
        "75/25 train/calibration partitioning; the nonconformity scores "
        "from every calibration fold are pooled and the (1−α)(n+1)/n "
        "quantile of each class is taken. Finally the script bundles "
        "everything — coefficients, scaler stats, quantiles, CEA "
        "rollback — into a single JSON.", indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 12 — Dashboards
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 12.  The interactive dashboards — HTML, CSS, JavaScript", level=1)
    add_para(doc,
        "The three dashboards on the companion site (calculator, savings, "
        "callgraph) are static HTML pages served by GitHub Pages. Each "
        "page bundles its own JavaScript so all computation happens in "
        "the user's browser — no data leaves the machine.", indent=True)

    add_heading(doc, "12.1  Anatomy of an HTML page", level=2)
    add_code_block(doc,
        '<!DOCTYPE html>\n'
        '<html lang="en">\n'
        '<head>\n'
        '  <meta charset="UTF-8">\n'
        '  <title>Risk calculator</title>\n'
        '  <style> /* CSS — visual styling lives here */ </style>\n'
        '</head>\n'
        '<body>\n'
        '  <header>...</header>\n'
        '  <main>...</main>\n'
        '  <script> /* JavaScript — behaviour lives here */ </script>\n'
        '</body>\n'
        '</html>\n', language="html")
    add_para(doc,
        "Three layers do three different jobs. HTML (HyperText Markup "
        "Language) defines the structure: headings, paragraphs, form "
        "inputs. CSS (Cascading Style Sheets) defines the visual style: "
        "colours, fonts, spacing. JavaScript defines the behaviour: what "
        "happens when the user types in a box, what computations to run, "
        "what to display.", indent=True)

    add_heading(doc, "12.2  How the risk calculator works", level=2)
    add_para(doc,
        "The risk calculator's JavaScript is roughly thirty lines of "
        "logic. On page load it fetches model_assets.json, extracts the "
        "Firth coefficients and the conformal quantiles, and attaches an "
        "event listener to every form input. Whenever the user changes "
        "any input, the recompute function fires: it reads all 21 "
        "feature values, z-scales them using the means and standard "
        "deviations from the JSON, computes the dot product with the "
        "coefficients, adds the intercept, and applies the sigmoid to "
        "obtain a probability. That probability is then used to compute "
        "the conformal prediction set at the user-selected α.", indent=True)
    calc_src = read_site("calculator.html")
    # Extract just the JS recompute function for illustration
    js_start = calc_src.find("function recompute")
    js_end = calc_src.find("\ndocument.addEventListener", js_start)
    add_code_block(doc, calc_src[js_start:js_end].rstrip(), language="javascript",
                     max_lines=70)
    add_para(doc,
        "Notice that the dot product is a single for loop. Each "
        "iteration takes one raw feature value, z-scales it by "
        "subtracting the training-time mean and dividing by the "
        "training-time SD, multiplies the result by the corresponding "
        "Firth coefficient, and accumulates it into z. After the loop, "
        "the intercept is added and the sigmoid applied: "
        "p = 1 / (1 + exp(-z)). That gives the predicted probability. "
        "The conformal logic that follows checks whether 1 − p is below "
        "each class-conditional quantile; the answer determines the "
        "prediction set and the recommendation.", indent=True)

    add_heading(doc, "12.3  How the savings calculator works", level=2)
    add_para(doc,
        "The savings calculator's JavaScript multiplies the per-patient "
        "ΔCost and ΔQALY values from model_assets.json by the user-"
        "specified cohort size, applies the chosen discount rate over "
        "the chosen horizon, and presents the result as KPI cards and "
        "a net-monetary-benefit bar chart. No external libraries are "
        "used — the bars are styled <div> elements whose width property "
        "is set proportionally.", indent=True)

    add_heading(doc, "12.4  How the interactive callgraph works", level=2)
    add_para(doc,
        "The callgraph dashboard uses one external library — vis-network — "
        "loaded from a CDN. On page load the script fetches "
        "callgraph.json (28 nodes, 27 edges), maps each script to a "
        "vis-network node with category-coloured fill, attaches the per-"
        "node function inventory as a hidden property, and renders the "
        "network with a force-directed layout. When the user clicks a "
        "node, the function-inventory is shown in a side panel. The "
        "physics toggle and edge-visibility toggle are simple "
        "checkbox-driven calls to setOptions and setData.", indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 13 — The callgraph
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 13.  The callgraph — how the scripts depend on each other", level=1)
    add_para(doc,
        "Every analysis script in the repository imports _shared.py. "
        "Many also import each other: 16_voi_evpi.py imports the strategy-"
        "rollback functions from 10_11_cea_pairwise.py; 26_main_figures.py "
        "and 29_main_figures_jnnp.py read CSV outputs that other scripts "
        "have produced; 27_build_jnnp_manuscript.py reads figures and "
        "results and assembles them into the Word file.", indent=True)
    add_para(doc,
        "The callgraph is generated by 28_make_callgraph.py (for the "
        "Markdown / Mermaid version embedded in CALLGRAPH.md) and "
        "31_export_callgraph_json.py (for the JSON consumed by the "
        "interactive dashboard). Both work by parsing each .py file's "
        "Abstract Syntax Tree (AST) — Python's machine-readable "
        "representation of its own source code — and extracting "
        "the import statements and the top-level function definitions.",
        indent=True)
    add_callout(doc, "How to use the callgraph when you change something.",
                 "Before you modify any function in _shared.py, look at "
                 "the callgraph and ask yourself: which scripts import "
                 "this function? Any script with an arrow pointing into "
                 "_shared.py depends on its behaviour. Test downstream "
                 "before publishing the change.")
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 14 — Other scripts in one paragraph each
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 14.  Other scripts — what they do, in one paragraph each", level=1)
    others = [
        ("03_dca.py",
         "Decision-curve analysis. Reads cached out-of-fold predictions, "
         "computes Vickers net benefit at every probability threshold "
         "from 0 to 30%, plots the model curve against 'treat all' and "
         "'treat none', and writes the summary CSV used by Figure 2 panel B."),
        ("05_temporal_leakage.py",
         "Temporal-leakage audit. Excludes the 24-hour and 48-hour rolling "
         "lab/vital features and the prophylactic-AED indicator from the "
         "eICU feature set, refits the model, and reports the resulting "
         "AUC. Used in Figure 1 panel C."),
        ("06_overfitting.py",
         "Overfitting diagnostics. Repeats the BalancedRandomForest fit "
         "with nested cross-validation, compares variable-importance "
         "ranks across folds (Spearman ρ across 25 folds = 0.98), and "
         "produces a learning-curve plot."),
        ("07_missing_data.py",
         "Missing-data sensitivity. Runs Little's MCAR test, computes the "
         "missingness-versus-outcome chi-square per feature, and pools "
         "AUC estimates across ten multiply-imputed datasets via Rubin's "
         "rules."),
        ("08_eicu_cohort.py",
         "eICU cohort definition sensitivity. Fits the model in each of "
         "four prespecified cohort strata (non-traumatic; non-traumatic + "
         "operative; APACHE-dx surgery/burr; traumatic-SDH negative "
         "control) and reports bootstrap 95% AUC CIs for each."),
        ("09_competing_risks.py",
         "Competing-risks survival analysis. Cause-specific Cox model "
         "with in-hospital death as the competing event, plus IPCW Fine–"
         "Gray subdistribution-hazard model (Geskus method). Reports "
         "concordance indices and the Grambsch–Therneau Schoenfeld test."),
        ("10_11_cea_pairwise.py",
         "Cost-effectiveness analysis with 10,000-iteration probabilistic "
         "sensitivity analysis. Defines the Params dataclass holding "
         "every probability, cost, and utility; samples each parameter "
         "from its own beta or gamma distribution per iteration; runs "
         "each of the four strategies; reports the cost-effectiveness "
         "plane and the acceptability curve."),
        ("12_nis_seizure_reclassify.py",
         "The NIS outcome-correction analysis. Distinguishes acute "
         "symptomatic seizure (R56.x, 780.39, G41.x) from pre-existing "
         "epilepsy (G40.x, 345.x) and reports both definitions' AUC. "
         "The collapse from 0.617 to 0.498 under the corrected outcome "
         "is the methodological-correction finding."),
        ("13_nis_grouped_lasso.py",
         "Group-LASSO with λ-path tuning on the NIS chronic+surgical "
         "cohort. Tests whether structured-sparsity regularisation can "
         "recover any signal under the corrected outcome; it cannot, "
         "reinforcing the outcome-correction finding."),
        ("15_radiology_nlp.py",
         "Regex-pattern radiology NLP extractor. Released as a tool in "
         "the repository (the script ships with a synthetic validation "
         "set) but not promoted to a paper contribution since we did not "
         "apply it to a real radiology corpus."),
        ("17_build_slides.py",
         "Builds the 15-minute oral-presentation deck using python-pptx. "
         "16 slides, journal-style palette, embedded figures and "
         "speaker-notes."),
        ("18_bidmc_optimize.py",
         "BIDMC model-optimization sweep. Runs Optuna for 40 trials each "
         "on XGBoost and LightGBM; fits a stacking ensemble; compares "
         "against the BalancedRandomForest baseline via DeLong tests. "
         "Confirms the AUC ceiling at 48 events."),
        ("19_transfer_learning.py",
         "eICU→BIDMC transfer-learning experiment. Trains an XGBoost on "
         "the seven shared features, scores BIDMC patients, and uses the "
         "score as an additional feature in the BIDMC model. Tests "
         "whether the transfer signal lifts BIDMC AUC; it does not."),
        ("20_build_manuscript.py",
         "An earlier manuscript-build script that produced the long-form "
         "draft used during revision. Superseded by 27_build_jnnp_"
         "manuscript.py which produces the journal-format submission."),
        ("21_imbalance_sweep.py",
         "Eleven-method class-imbalance sweep on BIDMC. Compares "
         "BalancedRF, class-weighted RF, six SMOTE-family oversamplers, "
         "XGBoost with scale_pos_weight, and XGBoost with focal loss. "
         "Confirms that none lift AUC beyond noise."),
        ("22_diverse_stacking.py",
         "Diverse-base stacking ensemble (logistic regression + Balanced"
         "RF + XGBoost + KNN + RBF-SVM with a logistic meta-learner) "
         "with optional isotonic recalibration."),
        ("23_tabpfn_eval.py",
         "Attempted TabPFN v2 evaluation. The script gracefully reports "
         "the API-credential requirement (TabPFN v2 is not free for "
         "research-grade use) and falls back without changing the "
         "headline results."),
        ("26_main_figures.py",
         "An earlier figure-builder that composes existing PNGs into "
         "composite plates via PIL. Superseded by 29_main_figures_jnnp.py "
         "which builds the figures natively in matplotlib for journal-"
         "style aesthetics."),
        ("28_make_callgraph.py",
         "Parses every .py file in scripts/ via Python's ast module, "
         "extracts imports and top-level function definitions, and "
         "emits a Markdown / Mermaid callgraph at CALLGRAPH.md."),
        ("31_export_callgraph_json.py",
         "The same parsing pipeline as 28 but emits the graph as JSON "
         "for the interactive callgraph dashboard to consume."),
        ("32_build_code_companion.py",
         "This script — the one that built the document you are reading."),
    ]
    for name, blurb in others:
        add_heading(doc, name, level=2)
        add_para(doc, blurb, indent=True)
    page_break(doc)

    # ─────────────────────────────────────────────────────────
    # Chapter 15 — Skills to make the code make sense
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "Chapter 15.  Further study — Claude skills that help "
                       "you learn the code", level=1)
    add_para(doc,
        "Claude (the AI assistant) ships with a library of 'skills' — "
        "task-focused expert prompts you can invoke by typing the skill "
        "name with a leading slash. The following skills are particularly "
        "relevant to studying this codebase. Each skill is invoked once "
        "and then guides the rest of your conversation.", indent=True)
    skill_recs = [
        ("/claude-code-guide",
         "When you want to ask Claude about Python idioms, scikit-learn "
         "patterns, or the Anthropic API itself. The most useful entry "
         "point for general programming questions."),
        ("/code-review-excellence",
         "When you want a structured second-pair-of-eyes review of a "
         "function or a script. Gives prioritised feedback: must-fix, "
         "should-fix, nice-to-have."),
        ("/documentation-writer",
         "When you want to add docstrings, README sections, or function-"
         "level comments. Guided by the Diátaxis framework (tutorial / "
         "how-to / reference / explanation)."),
        ("/python-project-structure",
         "When you want to understand why a Python project is organised "
         "the way it is. Covers public-API design, package layout, "
         "module boundaries."),
        ("/scientific-writing",
         "When you are translating an analysis result into manuscript "
         "prose. Covers paragraph structure, hedging, and the IMRAD "
         "convention."),
        ("/scientific-critical-thinking",
         "When you want to interrogate a methodological choice. Covers "
         "bias diagnostics, evidence-grading frameworks (GRADE, Cochrane "
         "Risk of Bias)."),
        ("/visualization-best-practices",
         "When you want to refine a figure. Covers chart-type selection, "
         "colour-blind safety, annotation placement, and grayscale "
         "fidelity."),
        ("/statistical-analysis",
         "When you want help choosing a statistical test. Guides test "
         "selection, assumption checking, and reporting in the manner "
         "expected by clinical journals."),
        ("/the-humanizer",
         "When you want to remove AI-pattern phrasing from a draft. "
         "Auto-detects content type and rewrites with channel-specific "
         "rules."),
        ("/peer-review",
         "When you want a structured manuscript review applying CONSORT "
         "or STROBE or TRIPOD checklists. Best for actual review-writing "
         "rather than evidence evaluation."),
        ("/literature-review",
         "When you want a comprehensive, systematically-searched "
         "literature review on a topic. Produces PDF and Markdown "
         "outputs with verified citations."),
        ("/exploratory-data-analysis",
         "When you have a new dataset (CSV, h5ad, parquet, BAM, etc.) "
         "and want a structured initial inspection with quality metrics."),
        ("/scikit-learn",
         "When you have a specific question about scikit-learn — how to "
         "build a pipeline, how to interpret a cross-validation result, "
         "how to debug a fit-time error."),
        ("/statsmodels",
         "When you need formal frequentist inference — OLS / GLM / "
         "mixed-effects / time-series — that scikit-learn does not "
         "provide. Covers diagnostic plots and statsmodels' API."),
        ("/pymc",
         "When you want to fit a Bayesian model (hierarchical / mixed-"
         "effects / time-series) with MCMC. Useful for the Bayesian-LR "
         "deepening that section 24 references."),
        ("/improve-codebase-architecture",
         "When you have read most of the code and want to refactor. "
         "Identifies opportunities to consolidate, reduce duplication, "
         "and increase testability."),
        ("/peer-review",
         "When you want a structured pre-submission review of the "
         "manuscript drafts under docs/. Applies clinical-journal "
         "reporting standards."),
    ]
    for name, body in skill_recs:
        add_heading(doc, name, level=2)
        add_para(doc, body, indent=True)
    add_para(doc, "")
    add_callout(doc, "A short reading order for first-time exploration.",
                "(1) Read Chapter 1 of this document to learn the Python "
                "constructs you will see. (2) Read _shared.py with "
                "Chapter 2 of this document open beside you. (3) Read "
                "02_calibration.py and 24_firth_bayes_lr.py with Chapters 3 "
                "and 6 open. (4) Open the calculator dashboard in your "
                "browser, then come back to Chapter 12 to see how the "
                "JavaScript implements the same model the Python scripts "
                "trained. By that point the rest of the repository will "
                "feel familiar.")

    doc.save(OUT_PATH)
    print(f"[OK] {OUT_PATH}")
    print(f"     size: {os.path.getsize(OUT_PATH)/1024:.1f} KB")


if __name__ == "__main__":
    build()
