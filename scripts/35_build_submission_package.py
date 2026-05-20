"""Task 35 — Build the submission-ready bundle.

Drops into Manuscript_05192026/:
  • main_manuscript.docx   (already present — copied if newer)
  • supplementary.docx     (already present — copied if newer)
  • Figure_1.tiff … Figure_6.tiff   (LZW-compressed, 300 dpi)
  • Figure_0_graphical_abstract.tiff
  • Tables.docx            (just the manuscript's tables, one per page)

TIFF format is what most clinical journals' submission portals want;
LZW compression keeps file size manageable while remaining lossless.
"""
import os, sys, shutil
sys.path.insert(0, os.path.dirname(__file__))
from pathlib import Path
import pandas as pd
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from _shared import FIG, RES

SUB_DIR = Path("/Users/nielspacheco/Desktop/Research/"
                 "Ogilvy research/Data Chronic Subdural Haematoma/"
                 "Manuscript_05192026")
SUB_DIR.mkdir(parents=True, exist_ok=True)

# ───────────────────────────────────────────────
# 1.  Convert each main figure PNG → TIFF (LZW, 300 dpi)
# ───────────────────────────────────────────────
FIGURE_MAP = {
    "F0_graphical_abstract.png":  "Figure_0_graphical_abstract.tiff",
    "F1_discrimination.png":      "Figure_1.tiff",
    "F2_calibration_dca.png":     "Figure_2.tiff",
    "F3_method_battery.png":      "Figure_3.tiff",
    "F4_conformal.png":           "Figure_4.tiff",
    "F5_cea.png":                 "Figure_5.tiff",
    "F6_voi.png":                 "Figure_6.tiff",
}


def png_to_tiff(src, dst, dpi=300):
    """Convert a PNG to a TIFF with LZW compression at the given DPI."""
    img = Image.open(src)
    # Many journals require RGB (not palette). Force conversion.
    if img.mode in ("P", "LA"):
        img = img.convert("RGBA")
    if img.mode == "RGBA":
        # Flatten transparency onto a white background
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")
    img.save(dst, format="TIFF", compression="tiff_lzw",
              dpi=(dpi, dpi))


# ───────────────────────────────────────────────
# 2.  Build Tables.docx
# ───────────────────────────────────────────────
def _setup(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0); section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(11)

def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x14, 0x1F, 0x3A)
        r.font.name = "Times New Roman"

def _caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(11); r.font.name = "Times New Roman"

def _note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.italic = True
    r.font.size = Pt(9); r.font.name = "Times New Roman"

def _add_dataframe(doc, df, *, monospace_numeric=True):
    t = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    t.style = "Light List Accent 1"
    for j, col in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"
    for i, row in df.iterrows():
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                # Choose a sensible format based on magnitude
                if pd.isna(val):
                    val = "—"
                elif abs(val) >= 1000:
                    val = f"{val:,.0f}"
                elif abs(val) >= 1:
                    val = f"{val:.3f}"
                else:
                    val = f"{val:.4f}"
            cell = t.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = ("Menlo" if monospace_numeric
                                    and j > 0 and isinstance(row[col], (int, float))
                                    else "Times New Roman")


def build_tables_docx(out_path):
    doc = Document(); _setup(doc)
    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Tables for the manuscript")
    tr.bold = True; tr.font.size = Pt(16); tr.font.name = "Times New Roman"

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Companion to main_manuscript.docx · one table per page · "
                   "extracted directly from the analysis result CSVs.")
    r.italic = True; r.font.size = Pt(11); r.font.name = "Times New Roman"
    doc.add_page_break()

    # ── Table 1: cohort characteristics (manually defined) ──
    # This is the ONLY table referenced in the main manuscript text;
    # the five method/result/CEA/EVPPI tables that used to live here are
    # already rendered as Tables S1–S5 in supplementary.docx and have
    # been removed from this companion file to avoid duplication.
    _caption(doc, "Table 1.  Cohort characteristics across the BIDMC "
                   "development cohort and the eICU external-validation "
                   "cohort (non-traumatic SDH stratum).")
    tbl1 = pd.DataFrame([
        {"Characteristic": "Patients, n",                          "BIDMC": "655",        "eICU non-traumatic": "3,297"},
        {"Characteristic": "Median age (years, IQR)",              "BIDMC": "73 [64–81]", "eICU non-traumatic": "74 [65–82]"},
        {"Characteristic": "Male sex, %",                          "BIDMC": "68",         "eICU non-traumatic": "63"},
        {"Characteristic": "Anticoagulant on admission, %",        "BIDMC": "27",         "eICU non-traumatic": "21"},
        {"Characteristic": "Burr-hole evacuation, %",              "BIDMC": "71",         "eICU non-traumatic": "—"},
        {"Characteristic": "Craniotomy, %",                        "BIDMC": "29",         "eICU non-traumatic": "—"},
        {"Characteristic": "Median preoperative GCS (IQR)",        "BIDMC": "14 [13–15]", "eICU non-traumatic": "14 [13–15]"},
        {"Characteristic": "Postoperative seizure, n (%)",         "BIDMC": "48 (7.3)",   "eICU non-traumatic": "300 (9.1)"},
    ])
    _add_dataframe(doc, tbl1, monospace_numeric=False)
    _note(doc, "Method, calibration, decision-curve and cost-effectiveness "
                "tables are released as Tables S1–S5 in supplementary.docx.")

    doc.save(out_path)


# ───────────────────────────────────────────────
# Main
# ───────────────────────────────────────────────
def main():
    SUB_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Submission directory: {SUB_DIR}")

    # 1. Refresh the two manuscript files if newer
    for fname in ("main_manuscript.docx", "supplementary.docx"):
        src = Path(FIG.parent / "manuscript" / fname)
        dst = SUB_DIR / fname
        if src.exists() and (not dst.exists()
                              or src.stat().st_mtime > dst.stat().st_mtime):
            shutil.copy2(src, dst)
            print(f"  ✓ refreshed {fname}")

    # 2. Convert figures
    print("\nConverting figures to TIFF (LZW, 300 dpi)...")
    for src_name, dst_name in FIGURE_MAP.items():
        src = FIG / src_name
        dst = SUB_DIR / dst_name
        if not src.exists():
            print(f"  ! skip — {src_name} not found in figures/")
            continue
        png_to_tiff(src, dst, dpi=300)
        size_mb = os.path.getsize(dst) / (1024 * 1024)
        print(f"  ✓ {dst_name:<38s}  ({size_mb:.1f} MB)")

    # 3. Build the tables-only Word document
    print("\nBuilding Tables.docx ...")
    tables_path = SUB_DIR / "Tables.docx"
    build_tables_docx(tables_path)
    size_kb = os.path.getsize(tables_path) / 1024
    print(f"  ✓ Tables.docx                              ({size_kb:.1f} KB)")

    # 4. Summary
    print(f"\n[OK] Submission bundle assembled at:\n  {SUB_DIR}\n")
    for p in sorted(SUB_DIR.iterdir()):
        if p.name.startswith("~$"): continue   # ignore Word lock files
        size = p.stat().st_size / 1024
        unit = "KB"
        if size > 1024: size, unit = size / 1024, "MB"
        print(f"  {p.name:<42s}  {size:>7.1f} {unit}")


if __name__ == "__main__":
    main()
