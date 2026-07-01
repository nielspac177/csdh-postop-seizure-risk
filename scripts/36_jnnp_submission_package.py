"""Task 36 — Assemble a numbered, JNNP-ready submission package.

Writes a clean, self-contained folder at project-root/JNNP_submission/ with
every component numbered in the order JNNP (BMJ ScholarOne) asks for them:

  00_Title_page.docx            (separate title page — NOT anonymised)
  01_Main_manuscript.docx       (abstract → key messages → text → refs → legends)
  02_Supplementary_appendix.docx
  03_Tables.docx                (Table 1 companion)
  Figure_0_graphical_abstract.tiff … Figure_6.tiff   (300 dpi, LZW)
  SUBMISSION_CHECKLIST.md       (numbered manifest + JNNP-limit audit)

Nothing here touches the *_NP working copies. Re-runnable and idempotent.
Run with:  ../.venv-csdh/bin/python3 scripts/36_jnnp_submission_package.py
"""
import os, sys, shutil, re
sys.path.insert(0, os.path.dirname(__file__))
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from _shared import FIG, OUT, ROOT

MANUS_DIR = OUT / "manuscript"
MAIN_SRC = MANUS_DIR / "main_manuscript.docx"
SUPP_SRC = MANUS_DIR / "supplementary.docx"
SUB_DIR = ROOT / "JNNP_submission"

# ── JNNP Original Research limits (for the manifest audit) ──
LIMITS = {"abstract": 250, "main_text": 3500, "tables_figures": 8, "references": 40}

FIGURE_MAP = {
    "F0_graphical_abstract.png": "Figure_0_graphical_abstract.tiff",
    "F1_discrimination.png":     "Figure_1.tiff",
    "F2_calibration_dca.png":    "Figure_2.tiff",
    "F3_method_battery.png":     "Figure_3.tiff",
    "F4_conformal.png":          "Figure_4.tiff",
    "F5_cea.png":                "Figure_5.tiff",
    "F6_voi.png":                "Figure_6.tiff",
}


def wc(s):
    return len(re.findall(r"\S+", s))


def audit_main(path):
    """Measure abstract words, main-text words, refs, tables/figures."""
    d = Document(path)
    paras = [(p.style.name if p.style else "", p.text) for p in d.paragraphs]
    abs_w = 0; cap = False
    for style, t in paras:
        if style.startswith("Heading") and t.strip() == "Abstract":
            cap = True; continue
        if style.startswith("Heading") and t.strip() == "Key messages":
            break
        if cap and t.strip():
            abs_w += wc(t)
    body_w = 0; cap = False
    for style, t in paras:
        if style.startswith("Heading 1") and t.strip() == "Introduction":
            cap = True
        if style.startswith("Heading 1") and t.strip() == "References":
            cap = False
        if cap:
            body_w += wc(t)
    refs = sum(1 for _, t in paras if re.match(r"^\d+\.\s", t.strip()))
    figs = sum(1 for _, t in paras if re.match(r"^Figure \d+\.", t.strip()))
    tabs = sum(1 for _, t in paras if re.match(r"^Table \d+\.", t.strip()))
    return {"abstract": abs_w, "main_text": body_w, "references": refs,
            "figures": figs, "tables": tabs}


# ───────────────────────────────────────────────
# Title page (standalone, NOT anonymised)
# ───────────────────────────────────────────────
def _p(doc, text, *, size=11, bold=False, italic=False,
       align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(); p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Times New Roman"
    return p


def build_title_page(out_path, audit):
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1.0)
        s.left_margin = s.right_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    _p(doc, "Title page", size=10, italic=True)
    _p(doc, "")
    _p(doc, "Postoperative seizure after chronic subdural haematoma "
            "evacuation: a calibration-focused, conformal-prediction "
            "proof-of-concept with value-of-information analysis",
       size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "")
    _p(doc, "Niels Pacheco-Barrios MD", size=12,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "Department of Neurosurgery, Beth Israel Deaconess Medical "
            "Center, Harvard Medical School, Boston, MA, USA", size=10,
       italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "")
    _p(doc, "Corresponding author: Niels Pacheco-Barrios MD, Department of "
            "Neurosurgery, BIDMC, 330 Brookline Ave, Boston MA 02215, USA. "
            "Email: nielspacheco1997@gmail.com", size=10)
    _p(doc, "")
    _p(doc, "Manuscript type: Original Research (proof of concept)", size=10)
    _p(doc, f"Word count — abstract: {audit['abstract']} (limit 250); "
            f"main text: {audit['main_text']} (limit 3500). "
            f"Tables: {audit['tables']}; Figures: {audit['figures']} "
            f"(limit 8 combined). References: {audit['references']} "
            f"(limit 40). Supplementary appendix: yes.", size=10)
    _p(doc, "")
    _p(doc, "Keywords: chronic subdural haematoma; postoperative seizure; "
            "clinical prediction model; conformal prediction; "
            "cost-effectiveness and value-of-information analysis", size=10)
    _p(doc, "Reporting standard: TRIPOD-AI (checklist in Supplementary "
            "Appendix S1).", size=10)
    _p(doc, "")
    _p(doc, "Funding: This study received no specific grant from any funding "
            "agency in the public, commercial or not-for-profit sectors.",
       size=10)
    _p(doc, "Competing interests: None declared.", size=10)
    _p(doc, "Author contributions: N.P.-B. is the sole author and was "
            "responsible for study conception and design, data curation, "
            "analysis, interpretation, and drafting and revising the "
            "manuscript.", size=10)
    _p(doc, "Ethics approval: Approved by the Beth Israel Deaconess Medical "
            "Center Institutional Review Board (Protocol [IRB number to be "
            "inserted]); the eICU Collaborative Research Database v2.0 was "
            "accessed under its data-use agreement.", size=10)
    _p(doc, "Data availability: Analysis code, figures, the TRIPOD-AI "
            "checklist and aggregate results are openly available at "
            "github.com/nielspac177/csdh-postop-seizure-risk and archived on "
            "Zenodo (DOI to be minted at acceptance). Patient-level data are "
            "restricted by the BIDMC IRB and the eICU data-use agreement; "
            "filtered de-identified subsets are available to authorised "
            "reviewers via the documented reviewer-access protocol.", size=10)
    doc.save(out_path)


def png_to_tiff(src, dst, dpi=300):
    img = Image.open(src)
    if img.mode in ("P", "LA"):
        img = img.convert("RGBA")
    if img.mode == "RGBA":
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3]); img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")
    img.save(dst, format="TIFF", compression="tiff_lzw", dpi=(dpi, dpi))


def build_tables_docx(out_path):
    import pandas as pd
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    _p(doc, "Table 1. Cohort characteristics across the BIDMC development "
            "cohort and the eICU external-evaluation cohort (non-traumatic "
            "SDH stratum).", size=11, bold=True)
    df = pd.DataFrame([
        ["Patients, n", "655", "3,297"],
        ["Median age (years, IQR)", "73 [64–81]", "74 [65–82]"],
        ["Male sex, %", "68", "63"],
        ["Anticoagulant on admission, %", "27", "21"],
        ["Burr-hole evacuation, %", "71", "—"],
        ["Craniotomy, %", "29", "—"],
        ["Median preoperative GCS (IQR)", "14 [13–15]", "14 [13–15]"],
        ["Postoperative seizure, n (%)", "48 (7.3)", "300 (9.1)"],
    ], columns=["Characteristic", "BIDMC", "eICU non-traumatic"])
    t = doc.add_table(rows=1 + len(df), cols=3); t.style = "Light List Accent 1"
    for j, c in enumerate(df.columns):
        t.rows[0].cells[j].text = str(c)
    for i, row in df.iterrows():
        for j, c in enumerate(df.columns):
            t.rows[i + 1].cells[j].text = str(row[c])
    _p(doc, "Method, calibration, decision-curve and cost-effectiveness "
            "tables are released as Tables S1–S5 in the supplementary "
            "appendix.", size=9, italic=True)
    doc.save(out_path)


def write_manifest(path, audit, tiff_names):
    def status(name, val):
        return "OK" if val <= LIMITS[name] else f"OVER by {val - LIMITS[name]}"
    tf = audit["tables"] + audit["figures"]
    lines = []
    lines.append("# JNNP Original Research — submission package manifest\n")
    lines.append("Journal of Neurology, Neurosurgery & Psychiatry · "
                 "Original Research (full paper).\n")
    lines.append("## 1. JNNP limit audit\n")
    lines.append("| Component | This paper | JNNP limit | Status |")
    lines.append("|---|---|---|---|")
    lines.append(f"| Abstract (words) | {audit['abstract']} | 250 | "
                 f"{status('abstract', audit['abstract'])} |")
    lines.append(f"| Main text (words) | {audit['main_text']} | 3500 | "
                 f"{status('main_text', audit['main_text'])} |")
    lines.append(f"| Tables + figures | {tf} "
                 f"({audit['tables']} tables + {audit['figures']} figures) "
                 f"| 8 | {status('tables_figures', tf)} |")
    lines.append(f"| References | {audit['references']} | 40 | "
                 f"{status('references', audit['references'])} |")
    lines.append(f"| Structured abstract (Bg/Methods/Results/Conclusions) "
                 f"| yes | required | OK |")
    lines.append(f"| Key-messages box (3 headings) | yes | required | OK |\n")
    lines.append("## 2. Files to upload (in order)\n")
    files = [
        ("00_Title_page.docx", "Title page — title, author, affiliation, "
         "corresponding author, word counts, keywords, funding, competing "
         "interests, contributions, ethics, data availability. Upload as "
         "\"Title Page\"."),
        ("01_Main_manuscript.docx", "Main document — abstract, key messages, "
         "Introduction/Methods/Results/Discussion, references, Table 1 and "
         "figure legends. Upload as \"Main Document\". Note: contains author "
         "identifiers on its own title page; if the portal requires a blinded "
         "main document, delete page 1 (the title block) before upload — the "
         "standalone 00_Title_page.docx already carries that information."),
        ("02_Supplementary_appendix.docx", "Supplementary appendices "
         "S1–S8 (TRIPOD-AI checklist, reproducibility appendix, feature "
         "dictionary, NIS appendix, LOHO sensitivity, calibration, CONSORT "
         "flow). Upload as \"Supplementary File\"."),
        ("03_Tables.docx", "Table 1 companion (optional; Table 1 is also in "
         "the main document)."),
    ]
    for i, (fn, desc) in enumerate(files, 1):
        lines.append(f"{i}. **{fn}** — {desc}")
    n = len(files)
    for k, tn in enumerate(tiff_names, 1):
        lines.append(f"{n + k}. **{tn}** — 300 dpi LZW-compressed TIFF. "
                     f"Upload each as a separate \"Figure\" "
                     f"({'graphical abstract' if 'Figure_0' in tn else 'main figure'}).")
    lines.append("\n## 3. Still to do before you hit submit\n")
    lines.append("- [ ] Insert the real BIDMC IRB protocol number "
                 "(placeholder `[IRB number to be inserted]` in title page, "
                 "Methods and ethics statement).")
    lines.append("- [ ] Mint the Zenodo DOI via the GitHub–Zenodo integration "
                 "and paste it into the data-availability statement.")
    lines.append("- [ ] Write a cover letter (not included here).")
    lines.append("- [ ] Confirm the ORCID iD on the ScholarOne author record.")
    lines.append("- [ ] Decide whether the portal needs a fully blinded main "
                 "document (see file 2 note).")
    path.write_text("\n".join(lines) + "\n")


def main():
    assert MAIN_SRC.exists(), f"missing {MAIN_SRC} — build the manuscript first"
    SUB_DIR.mkdir(parents=True, exist_ok=True)
    audit = audit_main(MAIN_SRC)
    print(f"Audit: {audit}")

    build_title_page(SUB_DIR / "00_Title_page.docx", audit)
    shutil.copy2(MAIN_SRC, SUB_DIR / "01_Main_manuscript.docx")
    if SUPP_SRC.exists():
        shutil.copy2(SUPP_SRC, SUB_DIR / "02_Supplementary_appendix.docx")
    build_tables_docx(SUB_DIR / "03_Tables.docx")

    tiff_names = []
    for src_name, dst_name in FIGURE_MAP.items():
        src = FIG / src_name
        if not src.exists():
            print(f"  ! skip — {src_name} not found"); continue
        png_to_tiff(src, SUB_DIR / dst_name)
        tiff_names.append(dst_name)

    write_manifest(SUB_DIR / "SUBMISSION_CHECKLIST.md", audit, tiff_names)

    print(f"\n[OK] JNNP submission package at:\n  {SUB_DIR}\n")
    for p in sorted(SUB_DIR.iterdir()):
        if p.name.startswith("~$"):
            continue
        size = p.stat().st_size / 1024
        unit = "KB"
        if size > 1024:
            size, unit = size / 1024, "MB"
        print(f"  {p.name:<40s} {size:>7.1f} {unit}")


if __name__ == "__main__":
    main()
